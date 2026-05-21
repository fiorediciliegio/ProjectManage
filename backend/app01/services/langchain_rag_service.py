from pathlib import Path
from django.conf import settings
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from docx import Document as DocxDocument
from openpyxl import load_workbook
import fitz
import uuid
import re
import requests
import statistics
import pdfplumber
import tempfile
from typing import List
from openai import OpenAI
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams, Filter, FieldCondition, MatchValue, MatchAny
from langchain_core.embeddings import Embeddings
from langchain_qdrant import QdrantVectorStore
from paddleocr import PaddleOCR
from app01.services.elasticsearch_service import (
    index_chunks_to_elasticsearch,
    delete_file_chunks_from_elasticsearch,
    keyword_search_file_chunks)

# ———————————————————— 文件预处理 ————————————————————
def build_base_metadata(file_obj):
    file_name = f'{file_obj.NAME_File}{file_obj.FORM_File or ""}'
    return {
        'file_id': file_obj.pk,
        'file_name': file_name,
        'project_id': file_obj.ID_Project_id,
        'project_name': file_obj.ID_Project.NAME_Project,
        'file_extension': (file_obj.FORM_File or '').lower(),
    }

# 结构化数据
def make_block(text, block_type, base_metadata, order, title_path=None, page=None, sheet_name=None, caption=None):
    return {
        **base_metadata,
        'text': (text or '').strip(),
        'block_type': block_type,
        'title_path': title_path or [],
        'page': page,
        'sheet_name': sheet_name,
        'caption': caption,
        'order': order,
    }

# 定义不同文件类型和 block 类型的切分策略
def get_semantic_merge_profile(block):
    file_extension = block.get('file_extension')
    block_type = block.get('block_type')

    if block_type in ['pdf_table', 'table', 'sheet_table']:
        return {
            'mergeable': False,
            'target_chars': 1600,
            'max_chars': 2200,
        }

    if block_type in ['title', 'pdf_title']:
        return {
            'mergeable': False,
            'target_chars': 0,
            'max_chars': 0,
        }

    if file_extension == '.pdf':
        return {
            'mergeable': block_type in ['pdf_paragraph', 'pdf_ocr_text', 'pdf_merged_paragraph'],
            'target_chars': 900,
            'max_chars': 1300,
        }

    if file_extension == '.docx':
        return {
            'mergeable': block_type in ['paragraph'],
            'target_chars': 900,
            'max_chars': 1300,
        }

    if file_extension in ['.txt', '.md', '.py', '.json', '.js', '.css', '.html']:
        return {
            'mergeable': True,
            'target_chars': 1000,
            'max_chars': 1400,
        }

    if file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.webp']:
        return {
            'mergeable': block_type in ['image_ocr_text'],
            'target_chars': 800,
            'max_chars': 1200,
        }

    return {
        'mergeable': True,
        'target_chars': 900,
        'max_chars': 1300,
    }

# 将文件切分为 block
def load_file_as_blocks(file_obj):
    file_path = file_obj.FILE.path
    file_extension = (file_obj.FORM_File or '').lower()
    base_metadata = build_base_metadata(file_obj)

    if file_extension in ['.txt', '.md', '.py', '.json', '.js', '.css', '.html']:
        return load_plain_file_blocks(file_path, base_metadata)

    if file_extension == '.docx':
        return load_docx_blocks(file_path, base_metadata)

    if file_extension in ['.xlsx', '.xlsm']:
        return load_xlsx_blocks(file_path, base_metadata)

    if file_extension == '.pdf':
        return load_pdf_blocks(file_path, base_metadata)

    if file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.webp']:
        return load_image_blocks(file_path, base_metadata)

    raise ValueError(f'暂不支持该文件类型入库：{file_extension}')

# 根据语义将 block 转 Document
def make_document_from_block(block):
    text = block.get('text', '').strip()
    title_path = block.get('title_path') or []
    caption = block.get('caption')

    prefix_lines = []

    if title_path:
        prefix_lines.append(f"标题路径：{' > '.join(title_path[-4:])}")

    if caption:
        prefix_lines.append(f"说明：{caption}")

    if block.get('block_type'):
        prefix_lines.append(f"内容类型：{block.get('block_type')}")

    page_content = '\n'.join(prefix_lines + [text]) if prefix_lines else text

    metadata = {
        key: value
        for key, value in block.items()
        if key != 'text'
    }

    return Document(
        page_content=page_content,
        metadata=metadata,
    )

# 根据语义合并 block
def blocks_to_semantic_documents(blocks):
    documents = []
    buffer_texts = []
    buffer_block = None
    buffer_chars = 0
    order = 0

    def flush_buffer():
        nonlocal buffer_texts, buffer_block, buffer_chars, order

        if not buffer_texts or not buffer_block:
            return

        merged_block = {
            **buffer_block,
            'text': '\n'.join(buffer_texts).strip(),
            'block_type': f"{buffer_block.get('block_type')}_semantic",
            'order': order,
        }

        documents.append(make_document_from_block(merged_block))

        buffer_texts = []
        buffer_block = None
        buffer_chars = 0
        order += 1

    for block in blocks:
        text = str(block.get('text') or '').strip()
        if not text:
            continue

        profile = get_semantic_merge_profile(block)
        block_type = block.get('block_type')
        title_path = block.get('title_path') or []
        page = block.get('page')
        sheet_name = block.get('sheet_name')

        if not profile['mergeable']:
            flush_buffer()
            single_block = {
                **block,
                'order': order,
            }
            documents.append(make_document_from_block(single_block))
            order += 1
            continue

        if buffer_block:
            buffer_profile = get_semantic_merge_profile(buffer_block)
            buffer_title_path = buffer_block.get('title_path') or []
            buffer_type = buffer_block.get('block_type')
            buffer_page = buffer_block.get('page')
            buffer_sheet_name = buffer_block.get('sheet_name')

            should_flush = (
                title_path != buffer_title_path
                or block_type != buffer_type
                or sheet_name != buffer_sheet_name
                or buffer_chars + len(text) > buffer_profile['max_chars']
            )

            if block.get('file_extension') == '.pdf':
                if page is not None and buffer_page is not None:
                    should_flush = should_flush or page - buffer_page > 1

            if should_flush:
                flush_buffer()

        if not buffer_block:
            buffer_block = block

        buffer_texts.append(text)
        buffer_chars += len(text)

        if buffer_chars >= profile['target_chars']:
            flush_buffer()

    flush_buffer()
    return merge_short_documents(documents)

# 短 document 二次合并
def merge_short_documents(documents, min_chars=220, max_chars=1200):
    merged = []
    buffer_doc = None

    def doc_key(doc):
        metadata = doc.metadata or {}
        return (
            metadata.get('file_id'),
            tuple(metadata.get('title_path') or []),
            metadata.get('block_type'),
            metadata.get('sheet_name'),
        )

    for doc in documents:
        content = doc.page_content or ''

        if buffer_doc is None:
            buffer_doc = doc
            continue

        buffer_content = buffer_doc.page_content or ''
        same_group = doc_key(buffer_doc) == doc_key(doc)

        if (
            same_group
            and len(buffer_content) < min_chars
            and len(buffer_content) + len(content) <= max_chars
        ):
            buffer_doc.page_content = buffer_content.rstrip() + '\n' + content.lstrip()
        else:
            merged.append(buffer_doc)
            buffer_doc = doc

    if buffer_doc is not None:
        merged.append(buffer_doc)

    return merged

# 将文件转化为 langchain 的 Document
def load_file_as_documents(file_obj):
    blocks = load_file_as_blocks(file_obj)
    return blocks_to_semantic_documents(blocks)

# 纯文本解析
def load_plain_file_blocks(file_path, base_metadata):
    text = Path(file_path).read_text(encoding='utf-8',errors='replace').strip()

    if not text:
        return []

    return [
        make_block(
            text=text,
            block_type='plain_text',
            base_metadata=base_metadata,
            order=0,
        )
    ]

# .docx 解析
def get_heading_level(paragraph):
    style_name = paragraph.style.name if paragraph.style else ''
    heading_map = {
        'Heading 1': 1,
        'Heading 2': 2,
        'Heading 3': 3,
        '标题 1': 1,
        '标题 2': 2,
        '标题 3': 3,
    }
    return heading_map.get(style_name)

def load_docx_blocks(file_path, base_metadata):
    document = DocxDocument(file_path)
    blocks = []
    title_stack = []
    order = 0

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        heading_level = get_heading_level(paragraph)
        if heading_level:
            title_stack = title_stack[:heading_level - 1]
            title_stack.append(text)
            blocks.append(
                make_block(
                    text=text,
                    block_type='title',
                    base_metadata=base_metadata,
                    order=order,
                    title_path=title_stack.copy(),
                )
            )
        else:
            blocks.append(
                make_block(
                    text=text,
                    block_type='paragraph',
                    base_metadata=base_metadata,
                    order=order,
                    title_path=title_stack.copy(),
                )
            )
        order += 1
    for table_index, table in enumerate(document.tables):
        table_text = convert_docx_table_to_markdown(table)
        if not table_text.strip():
            continue
        blocks.append(
            make_block(
                text=table_text,
                block_type='table',
                base_metadata=base_metadata,
                order=order,
                title_path=title_stack.copy(),
                caption=f'Word 表格 {table_index + 1}',
            )
        )
        order += 1
    return blocks

# .docx 表格转 Markdown
def convert_docx_table_to_markdown(table):
    rows = []

    for row in table.rows:
        cells = [
            cell.text.strip().replace('\n', ' ')
            for cell in row.cells
        ]
        if any(cells):
            rows.append(cells)

    if not rows:
        return ''

    max_cols = max(len(row) for row in rows)
    normalized_rows = [
        row + [''] * (max_cols - len(row))
        for row in rows
    ]

    header = normalized_rows[0]
    body = normalized_rows[1:]

    lines = []
    lines.append('| ' + ' | '.join(header) + ' |')
    lines.append('| ' + ' | '.join(['---'] * max_cols) + ' |')

    for row in body:
        lines.append('| ' + ' | '.join(row) + ' |')

    return '\n'.join(lines)

# .xlsx 解析
def load_xlsx_blocks(file_path, base_metadata):
    value_workbook = load_workbook(file_path, data_only=True)
    formula_workbook = load_workbook(file_path, data_only=False)
    blocks = []
    order = 0
    for sheet_name in value_workbook.sheetnames:
        value_sheet = value_workbook[sheet_name]
        formula_sheet = formula_workbook[sheet_name]

        table_text = convert_xlsx_sheet_to_markdown_with_formulas(
            value_sheet=value_sheet,
            formula_sheet=formula_sheet,
        )
        if not table_text.strip():
            continue
        blocks.append(
            make_block(
                text=table_text,
                block_type='sheet_table',
                base_metadata=base_metadata,
                order=order,
                sheet_name=sheet_name,
                caption=f'工作表：{sheet_name}',
            )
        )
        order += 1
    return blocks

# 表格转换
def convert_xlsx_sheet_to_markdown_with_formulas(value_sheet, formula_sheet):
    rows = []
    max_row = max(value_sheet.max_row, formula_sheet.max_row)
    max_column = max(value_sheet.max_column, formula_sheet.max_column)

    for row_index in range(1, max_row + 1):
        row_cells = []

        for column_index in range(1, max_column + 1):
            value_cell = value_sheet.cell(row=row_index, column=column_index)
            formula_cell = formula_sheet.cell(row=row_index, column=column_index)

            display_text = format_excel_cell_value_with_formula(
                value=formula_cell.value,
                calculated_value=value_cell.value,
            )

            row_cells.append(display_text)

        if any(cell.strip() for cell in row_cells):
            rows.append(row_cells)

    if not rows:
        return ''

    max_cols = max(len(row) for row in rows)

    normalized_rows = [
        row + [''] * (max_cols - len(row))
        for row in rows
    ]

    header = normalized_rows[0]
    body = normalized_rows[1:]

    lines = []
    lines.append('| ' + ' | '.join(header) + ' |')
    lines.append('| ' + ' | '.join(['---'] * max_cols) + ' |')

    for row in body:
        lines.append('| ' + ' | '.join(row) + ' |')
    return '\n'.join(lines)

# 单元格格式化
def format_excel_cell_value_with_formula(value, calculated_value):
    if value is None and calculated_value is None:
        return ''

    value_text = '' if value is None else str(value).strip()
    calculated_text = '' if calculated_value is None else str(calculated_value).strip()

    if value_text.startswith('='):
        if calculated_text:
            return f'{value_text}（计算结果：{calculated_text}）'
        return value_text

    return value_text

# .pdf 解析
# 添加页眉页脚识别函数
PDF_HEADER_RATIO = 0.045
PDF_FOOTER_RATIO = 0.055
PDF_REPEAT_RATIO = 0.6

def normalize_pdf_margin_text(text):
    return re.sub(r'\s+', '', text or '').strip()

def is_page_number_text(text):
    normalized = normalize_pdf_margin_text(text)

    return bool(
        re.match(r'^\d+$', normalized)
        or re.match(r'^第?\d+页$', normalized)
        or re.match(r'^[-—]\d+[-—]$', normalized)
        or re.match(r'^\d+/\d+$', normalized)
    )

def extract_text_from_pdf_block(block):
    lines = []
    max_font_size = 0
    font_sizes = []

    for line in block.get('lines', []):
        line_text_parts = []

        for span in line.get('spans', []):
            text = span.get('text', '').strip()
            if not text:
                continue

            line_text_parts.append(text)

            size = span.get('size', 0)
            if size:
                font_sizes.append(size)
                max_font_size = max(max_font_size, size)

        if line_text_parts:
            lines.append(''.join(line_text_parts))

    return '\n'.join(lines).strip(), max_font_size, font_sizes

# 统计重复页眉页脚文本
def collect_repeated_pdf_margin_texts(pdf):
    margin_text_counts = {}
    page_count = pdf.page_count

    if page_count < 3:
        return set()

    for page_index in range(page_count):
        page = pdf.load_page(page_index)
        page_height = page.rect.height
        page_dict = page.get_text('dict')

        for block in page_dict.get('blocks', []):
            if block.get('type') != 0:
                continue

            bbox = block.get('bbox')
            if not bbox:
                continue

            x0, y0, x1, y1 = bbox
            in_header = y1 <= page_height * PDF_HEADER_RATIO
            in_footer = y0 >= page_height * (1 - PDF_FOOTER_RATIO)

            if not in_header and not in_footer:
                continue

            text, _, _ = extract_text_from_pdf_block(block)
            normalized = normalize_pdf_margin_text(text)

            if not normalized:
                continue

            if len(normalized) > 80:
                continue

            margin_text_counts.setdefault(normalized, set()).add(page_index)

    min_repeat_pages = max(3, int(page_count * PDF_REPEAT_RATIO))

    return {
        text
        for text, page_indexes in margin_text_counts.items()
        if len(page_indexes) >= min_repeat_pages
    }

# 判断某个block是否应过滤
def should_skip_pdf_margin_block(block, page, repeated_margin_texts):
    bbox = block.get('bbox')
    if not bbox:
        return False

    x0, y0, x1, y1 = bbox
    page_height = page.rect.height

    in_header = y1 <= page_height * PDF_HEADER_RATIO
    in_footer = y0 >= page_height * (1 - PDF_FOOTER_RATIO)

    if not in_header and not in_footer:
        return False

    text, _, _ = extract_text_from_pdf_block(block)
    normalized = normalize_pdf_margin_text(text)

    if not normalized:
        return True
    if is_page_number_text(normalized):
        return True
    if normalized in repeated_margin_texts:
        return True
    return False

# 标题判断函数
def get_pdf_page_average_font_size(page_dict):
    font_sizes = []

    for block in page_dict.get('blocks', []):
        if block.get('type') != 0:
            continue

        for line in block.get('lines', []):
            for span in line.get('spans', []):
                size = span.get('size', 0)
                text = span.get('text', '').strip()

                if size and text:
                    font_sizes.append(size)

    if not font_sizes:
        return 12
    return statistics.median(font_sizes)

def is_pdf_title_block(text, max_font_size, average_font_size):
    stripped = (text or '').strip()

    if not stripped:
        return False

    normalized = stripped.replace('\n', '')

    heading_patterns = [
        r'^摘要$',
        r'^关键词',
        r'^Abstract$',
        r'^Keywords',
        r'^引言$',
        r'^结论$',
        r'^参考文献$',
        r'^第[一二三四五六七八九十\d]+章',
        r'^\d+(\.\d+)*\s+.+',
    ]

    if any(re.match(pattern, normalized, re.IGNORECASE) for pattern in heading_patterns):
        return True
    if max_font_size >= average_font_size + 2 and len(normalized) <= 80:
        return True
    return False

# 多栏判断函数
def is_two_column_page(text_blocks, page_width):
    if len(text_blocks) < 8:
        return False

    left_count = 0
    right_count = 0
    middle_count = 0

    for block in text_blocks:
        x0, y0, x1, y1 = block.get('bbox')
        center_x = (x0 + x1) / 2

        if center_x < page_width * 0.42:
            left_count += 1
        elif center_x > page_width * 0.58:
            right_count += 1
        else:
            middle_count += 1

    total = len(text_blocks)

    return left_count >= total * 0.3 and right_count >= total * 0.3 and middle_count <= total * 0.25

# 排序函数
def sort_pdf_text_blocks_by_layout(text_blocks, page_width):
    if not is_two_column_page(text_blocks, page_width):
        return sorted(
            text_blocks,
            key=lambda item: (
                round(item.get('bbox')[1], 1),
                round(item.get('bbox')[0], 1),
            )
        )
    left_blocks = []
    right_blocks = []
    full_width_blocks = []

    for block in text_blocks:
        x0, y0, x1, y1 = block.get('bbox')
        block_width = x1 - x0
        center_x = (x0 + x1) / 2

        if block_width >= page_width * 0.65:
            full_width_blocks.append(block)
        elif center_x < page_width / 2:
            left_blocks.append(block)
        else:
            right_blocks.append(block)

    full_width_blocks.sort(key=lambda item: (item.get('bbox')[1], item.get('bbox')[0]))
    left_blocks.sort(key=lambda item: (item.get('bbox')[1], item.get('bbox')[0]))
    right_blocks.sort(key=lambda item: (item.get('bbox')[1], item.get('bbox')[0]))

    return full_width_blocks + left_blocks + right_blocks

# .pdf 表格转Markdown
def convert_pdf_table_to_markdown(table):
    if not table:
        return ''

    rows = []

    for row in table:
        cleaned_row = [
            normalize_pdf_table_cell(cell)
            for cell in row
        ]

        if any(cleaned_row):
            rows.append(cleaned_row)

    if not rows:
        return ''

    max_cols = max(len(row) for row in rows)

    normalized_rows = [
        row + [''] * (max_cols - len(row))
        for row in rows
    ]

    header = normalized_rows[0]
    body = normalized_rows[1:]

    lines = []
    lines.append('| ' + ' | '.join(header) + ' |')
    lines.append('| ' + ' | '.join(['---'] * max_cols) + ' |')

    for row in body:
        lines.append('| ' + ' | '.join(row) + ' |')
    return '\n'.join(lines)

def normalize_pdf_table_cell(cell):
    if cell is None:
        return ''

    text = str(cell).strip()
    text = re.sub(r'\s+', ' ', text)
    text = text.replace('|', '/')

    return text

# .pdf 表格提取函数
def extract_pdf_table_blocks(file_path, base_metadata, start_order=0):
    blocks = []
    order = start_order

    with pdfplumber.open(file_path) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):
            tables = page.extract_tables()

            for table_index, table in enumerate(tables, start=1):
                table_text = convert_pdf_table_to_markdown(table)

                if not table_text.strip():
                    continue

                blocks.append(
                    make_block(
                        text=table_text,
                        block_type='pdf_table',
                        base_metadata=base_metadata,
                        order=order,
                        page=page_index,
                        caption=f'PDF 第 {page_index} 页表格 {table_index}',
                    )
                )
                order += 1
    return blocks

# 扫描页判断
def is_scanned_pdf_page(page, min_text_length=40):
    text = page.get_text('text').strip()
    images = page.get_images(full=True)

    return len(text) < min_text_length and len(images) > 0

_ocr_engine = None

def get_ocr_engine():
    global _ocr_engine

    if PaddleOCR is None:
        return None

    if _ocr_engine is None:
        _ocr_engine = PaddleOCR(
            lang='ch',
            text_detection_model_name='PP-OCRv5_mobile_det',
            text_recognition_model_name='PP-OCRv5_mobile_rec',
            use_doc_orientation_classify=False,
            use_doc_unwarping=False,
            use_textline_orientation=True,
            enable_mkldnn=False,
        )
    return _ocr_engine

# 把 pdf 页面渲染成 png
def render_pdf_page_to_image(page, zoom=2):
    matrix = fitz.Matrix(zoom, zoom)
    pixmap = page.get_pixmap(matrix=matrix, alpha=False)
    return pixmap.tobytes('png')

# OCR 结果解析函数
def extract_text_from_ocr_result(result, min_confidence=0.6):
    lines = []

    for page_result in result or []:
        if isinstance(page_result, dict):
            texts = page_result.get('rec_texts') or []
            scores = page_result.get('rec_scores') or []

            for index, text in enumerate(texts):
                score = scores[index] if index < len(scores) else 1
                if score >= min_confidence and str(text).strip():
                    lines.append(str(text).strip())
            continue

        if hasattr(page_result, 'json'):
            data = page_result.json
            texts = data.get('rec_texts') or []
            scores = data.get('rec_scores') or []

            for index, text in enumerate(texts):
                score = scores[index] if index < len(scores) else 1
                if score >= min_confidence and str(text).strip():
                    lines.append(str(text).strip())

            continue
        # 兼容旧版 PaddleOCR 返回格式
        for item in page_result or []:
            try:
                text = item[1][0]
                score = item[1][1]
            except (IndexError, TypeError):
                continue

            if score >= min_confidence and str(text).strip():
                lines.append(str(text).strip())
    return '\n'.join(lines)

# 单页 OCR 函数
def ocr_pdf_page(page):
    ocr_engine = get_ocr_engine()

    if ocr_engine is None:
        return ''

    image_bytes = render_pdf_page_to_image(page, zoom=2)

    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
        temp_file.write(image_bytes)
        temp_image_path = temp_file.name

    try:
        result = ocr_engine.ocr(temp_image_path)
    finally:
        Path(temp_image_path).unlink(missing_ok=True)

    return extract_text_from_ocr_result(result)


def load_pdf_blocks(file_path, base_metadata):
    pdf = fitz.open(file_path)
    blocks = []
    order = 0
    try:
        repeated_margin_texts = collect_repeated_pdf_margin_texts(pdf)

        for page_index in range(pdf.page_count):
            page = pdf.load_page(page_index)
            if is_scanned_pdf_page(page):
                ocr_text = ocr_pdf_page(page)
                if ocr_text.strip():
                    blocks.append(
                        make_block(
                            text=ocr_text,
                            block_type='pdf_ocr_text',
                            base_metadata=base_metadata,
                            order=order,
                            page=page_index + 1,
                            caption=f'PDF 第 {page_index + 1} 页 OCR 识别文本',
                        )
                    )
                    order += 1
                continue
            page_dict = page.get_text('dict')
            average_font_size = get_pdf_page_average_font_size(page_dict)
            raw_text_blocks = []

            for block in page_dict.get('blocks', []):
                if block.get('type') != 0:
                    continue
                if should_skip_pdf_margin_block(block, page, repeated_margin_texts):
                    continue
                bbox = block.get('bbox')
                if not bbox:
                    continue
                text, max_font_size, _ = extract_text_from_pdf_block(block)
                if not text:
                    continue
                raw_text_blocks.append({
                    'raw_block': block,
                    'bbox': bbox,
                    'text': text,
                    'max_font_size': max_font_size,
                })
            sorted_text_blocks = sort_pdf_text_blocks_by_layout(
                raw_text_blocks,
                page.rect.width,
            )
            for item in sorted_text_blocks:
                text = item['text']
                max_font_size = item['max_font_size']

                block_type = 'pdf_paragraph'
                if is_pdf_title_block(text, max_font_size, average_font_size):
                    block_type = 'pdf_title'

                blocks.append(
                    make_block(
                        text=text,
                        block_type=block_type,
                        base_metadata=base_metadata,
                        order=order,
                        page=page_index + 1,
                    )
                )
                order += 1
    finally:
        pdf.close()

    table_blocks = extract_pdf_table_blocks(
        file_path=file_path,
        base_metadata=base_metadata,
        start_order=order,
    )
    blocks.extend(table_blocks)
    blocks.sort(key=lambda item: (
        item.get('page') or 0,
        item.get('order') or 0,
    ))
    return blocks

# 图片解析
def load_image_blocks(file_path, base_metadata):
    ocr_engine = get_ocr_engine()

    if ocr_engine is None:
        raise ValueError('当前环境未安装 PaddleOCR，无法识别图片文字')

    result = ocr_engine.ocr(file_path)
    text = extract_text_from_ocr_result(result)

    if not text.strip():
        raise ValueError('图片中未识别到可入库文字')

    return [
        make_block(
            text=text,
            block_type='image_ocr_text',
            base_metadata=base_metadata,
            order=0,
            caption='图片 OCR 识别文本',
        )
    ]

# 切分策略函数
def get_split_profile(document):
    metadata = document.metadata or {}
    file_extension = metadata.get('file_extension')
    block_type = metadata.get('block_type')

    if block_type in ['pdf_table', 'table', 'sheet_table']:
        return {
            'chunk_size': 1800,
            'chunk_overlap': 80,
            'protect_limit': 2200,
        }

    if file_extension == '.pdf':
        return {
            'chunk_size': 1100,
            'chunk_overlap': 160,
            'protect_limit': None,
        }

    if file_extension == '.docx':
        return {
            'chunk_size': 1000,
            'chunk_overlap': 140,
            'protect_limit': None,
        }

    if file_extension in ['.xlsx', '.xlsm']:
        return {
            'chunk_size': 1600,
            'chunk_overlap': 80,
            'protect_limit': 2200,
        }

    if file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.webp']:
        return {
            'chunk_size': 800,
            'chunk_overlap': 100,
            'protect_limit': None,
        }

    return {
        'chunk_size': 1000,
        'chunk_overlap': 140,
        'protect_limit': None,
    }

# document 太长时兜底切分
def split_documents(documents):
    chunks = []
    for document in documents:
        content = document.page_content or ''
        profile = get_split_profile(document)
        protect_limit = profile.get('protect_limit')
        if protect_limit and len(content) <= protect_limit:
            chunks.append(document)
            continue

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=profile['chunk_size'],
            chunk_overlap=profile['chunk_overlap'],
            separators=['\n\n', '\n', '。', '；', '，', ' ', ''],
        )
        chunks.extend(splitter.split_documents([document]))
    for index, chunk in enumerate(chunks):
        chunk.metadata['chunk_index'] = index
    return chunks

# 测试入口函数
def preview_file_chunks(file_obj):
    documents = load_file_as_documents(file_obj)
    chunks = split_documents(documents)

    return {
        'documents_count': len(documents),
        'chunks_count': len(chunks),
        'chunks': [
            {
                'text': chunk.page_content[:300],
                'metadata': chunk.metadata,
            }
            for chunk in chunks[:5]
        ],
    }
# ———————————————————— embedding ————————————————————
# 本地 Qwen3 Embedding 适配器
class LocalOpenAICompatibleEmbeddings(Embeddings):
    def __init__(self):
        self.client = OpenAI(api_key='local', base_url=settings.EMBEDDING_BASE_URL)
        self.model = settings.EMBEDDING_MODEL

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        cleaned_texts = [
            self._clean_text(text)
            for text in texts
            if self._clean_text(text)
        ]

        if not cleaned_texts:
            return []

        response = self.client.embeddings.create(
            model=self.model,
            input=cleaned_texts,
        )

        return [item.embedding for item in response.data]

    def embed_query(self, text: str) -> List[float]:
        vectors = self.embed_documents([text])
        if not vectors:
            return []
        return vectors[0]

    @staticmethod
    def _clean_text(text):
        return (text or '').replace('\n', ' ').strip()

# Qdrant collection 初始化函数
def get_langchain_qdrant_client():
    return QdrantClient(
        url=settings.QDRANT_URL,
        prefer_grpc=False,
        timeout=10,
    )

def ensure_langchain_collection():
    client = get_langchain_qdrant_client()

    collection_names = [
        collection.name
        for collection in client.get_collections().collections
    ]

    if settings.LANGCHAIN_QDRANT_COLLECTION not in collection_names:
        client.create_collection(
            collection_name=settings.LANGCHAIN_QDRANT_COLLECTION,
            vectors_config=VectorParams(
                size=settings.QDRANT_VECTOR_SIZE,
                distance=Distance.COSINE,
            ),
        )
    return client

# LangChain VectorStore 获取函数
def get_langchain_vector_store():
    client = ensure_langchain_collection()
    embeddings = LocalOpenAICompatibleEmbeddings()

    return QdrantVectorStore(
        client=client,
        collection_name=settings.LANGCHAIN_QDRANT_COLLECTION,
        embedding=embeddings,
    )

# 删除旧向量
def delete_langchain_file_vectors(file_id):
    client = ensure_langchain_collection()

    client.delete(
        collection_name=settings.LANGCHAIN_QDRANT_COLLECTION,
        points_selector=Filter(
            must=[
                FieldCondition(
                    key='metadata.file_id',
                    match=MatchValue(value=file_id),
                )
            ]
        ),
    )
    delete_file_chunks_from_elasticsearch(file_id)

# 查询是否入库
def get_indexed_file_ids(project_id):
    ensure_langchain_collection()
    client = get_langchain_qdrant_client()
    qdrant_filter = Filter(
        must=[
            FieldCondition(
                key='metadata.project_id',
                match=MatchValue(value=project_id),
            )
        ]
    )

    file_ids = set()
    offset = None

    while True:
        points, offset = client.scroll(
            collection_name=settings.LANGCHAIN_QDRANT_COLLECTION,
            scroll_filter=qdrant_filter,
            limit=100,
            offset=offset,
            with_payload=True,
            with_vectors=False,
        )

        for point in points:
            metadata = (point.payload or {}).get('metadata') or {}
            file_id = metadata.get('file_id')
            if file_id is not None:
                file_ids.add(file_id)

        if offset is None:
            break
    return file_ids

# 写入 Qdrant
def index_file_to_qdrant_langchain(file_obj):
    documents = load_file_as_documents(file_obj)
    chunks = split_documents(documents)

    if not chunks:
        return {
            'file_id': file_obj.pk,
            'chunks_count': 0,
            'message': '文件中没有可入库的文本内容',
        }

    delete_langchain_file_vectors(file_obj.pk)

    vector_store = get_langchain_vector_store()

    ids = [
        str(uuid.uuid5(uuid.NAMESPACE_URL, f'file-{file_obj.pk}-chunk-{index}'))
        for index, _ in enumerate(chunks)
    ]

    vector_store.add_documents(
        documents=chunks,
        ids=ids,
    )
    elasticsearch_result = index_chunks_to_elasticsearch(file_obj, chunks)
    file_name = f'{file_obj.NAME_File}{file_obj.FORM_File or ""}'
    return {
        'file_id': file_obj.pk,
        'file_name': file_name,
        'chunks_count': len(chunks),
        'elasticsearch_chunks_count': elasticsearch_result.get('indexed_count', 0),
        'message': '文件向量和关键词索引入库成功',
    }
# ———————————————————— 混合检索 ————————————————————
# 检索函数
def search_file_chunks_langchain(question, project_id=None, limit=5):
    vector_store = get_langchain_vector_store()
    qdrant_filter = None
    if project_id is not None:
        qdrant_filter = Filter(
            must=[
                FieldCondition(
                    key='metadata.project_id',
                    match=MatchValue(value=project_id),
                )
            ]
        )

    docs_with_scores = vector_store.similarity_search_with_score(
        query=question,
        k=limit,
        filter=qdrant_filter,
    )

    results = []
    for doc, score in docs_with_scores:
        metadata = doc.metadata or {}
        results.append({
            'score': score,
            'file_id': metadata.get('file_id'),
            'project_id': metadata.get('project_id'),
            'project_name': metadata.get('project_name'),
            'file_name': metadata.get('file_name'),
            'file_extension': metadata.get('file_extension'),
            'chunk_index': metadata.get('chunk_index'),
            'block_type': metadata.get('block_type'),
            'page': metadata.get('page'),
            'sheet_name': metadata.get('sheet_name'),
            'text': doc.page_content,
        })
    return results

# 统一key判断Qdrant和ES搜到的是不是同一个chunk
def make_search_result_key(item):
    return (
        item.get('file_id'),
        item.get('chunk_index'),
    )

# RRF 融合函数
def reciprocal_rank_fusion(result_groups, rrf_k=60, top_k=8):
    fused_map = {}

    for group_name, results in result_groups:
        for rank, item in enumerate(results, start=1):
            key = make_search_result_key(item)

            if key in [(None, None), None]:
                continue

            if key not in fused_map:
                fused_map[key] = {
                    **item,
                    'rrf_score': 0,
                    'retrieval_sources': [],        # 记录 chunk 从哪里来，在vector第几名，在keyword第几名
                }

            fused_map[key]['rrf_score'] += 1 / (rrf_k + rank)
            fused_map[key]['retrieval_sources'].append({
                'source': group_name,
                'rank': rank,
                'score': item.get('score'),
            })

            if len(str(item.get('text') or '')) > len(str(fused_map[key].get('text') or '')):
                fused_map[key]['text'] = item.get('text')

    fused_results = list(fused_map.values())

    fused_results.sort(
        key=lambda item: item.get('rrf_score', 0),
        reverse=True,
    )

    return fused_results[:top_k]

# 混合检索函数
def hybrid_search_file_chunks(question, project_id=None, final_limit=8):
    vector_results = search_file_chunks_langchain(
        question=question,
        project_id=project_id,
        limit=25,               # Qdrant 召回前 25 个 chunk
    )
    keyword_results = keyword_search_file_chunks(
        query=question,
        project_id=project_id,
        limit=25,               # Elasticsearch 也先召回前 25 个 chunk
    )
    fused_results = reciprocal_rank_fusion(
        result_groups=[('vector', vector_results), ('keyword', keyword_results)],
        rrf_k=60,               # RRF算法将二者结果按排名融合排序，公式：rrf_score += 1 / (rrf_k + rank)
        top_k=RERANK_CANDIDATE_LIMIT,           # 再从融合结果里取 RERANK_CANDIDATE_LIMIT 个 chunk
    )
    rule_reranked_results = rule_rerank_search_results(
        search_results=fused_results,
        question=question,
        top_k=MODEL_RERANK_CANDIDATE_LIMIT,     # 规则 rerank
    )
    return rerank_search_results_with_model(
        question=question,
        search_results=rule_reranked_results,
        top_k=final_limit,                      # 使用模型进行二次rerank
    )
# ———————————————————— rerank ————————————————————
# 从用户问题中提取有检索价值的关键词
def extract_query_terms(question):
    text = str(question or '').strip()
    if not text:
        return []

    terms = re.findall(r'[A-Za-z][A-Za-z0-9\-_/.]*|[\u4e00-\u9fff]{2,}|\d+(?:\.\d+)?', text)

    stop_words = {'什么', '哪些', '怎么', '如何', '为什么', '这个', '那个','一下', '进行', '可以', '是否', '影响', '情况'}

    filtered_terms = []
    for term in terms:
        if term in stop_words:
            continue
        if len(term) <= 1:
            continue
        filtered_terms.append(term)
    return filtered_terms

# 判断是否属于解释型问题
def is_explanation_question(question):
    keywords = ['是什么', '为什么', '原理', '机制', '作用', '原因', '说明', '解释']
    text = str(question or '')
    return any(keyword in text for keyword in keywords)

# 判断是否是表格/数据型问题
def is_table_or_data_question(question):
    keywords = ['表', '参数', '数值', '编号', '型号', '尺寸', '金额','日期', '数量', '对比', '统计', '比例', '强度', '荷载']
    text = str(question or '')
    return any(keyword in text for keyword in keywords)

#判断是否属于参考文献
def is_reference_like_chunk(item):
    title_path = str(item.get('title_path') or '')
    text = str(item.get('text') or '')
    reference_markers = ['参考文献', 'References', '[J]', '[D]', '[S]', '出版社', '学报', '期刊']
    return any(marker in title_path or marker in text[:200] for marker in reference_markers)

# 获取检索来源名称
def get_retrieval_source_names(item):
    return {
        source.get('source')
        for source in item.get('retrieval_sources', [])
        if source.get('source')
    }

# 规则 rerank 参数
RERANK_CANDIDATE_LIMIT = 18
RERANK_FINAL_LIMIT = 8
RERANK_DOUBLE_HIT_BONUS = 1.20
RERANK_TITLE_MATCH_BONUS = 1.15
RERANK_GOOD_LENGTH_BONUS = 1.10
RERANK_SHORT_TEXT_PENALTY = 0.60
RERANK_LONG_TEXT_PENALTY = 0.90
RERANK_REFERENCE_PENALTY = 0.50
RERANK_TABLE_QUERY_BONUS = 1.18
RERANK_TABLE_MISMATCH_PENALTY = 0.90

# 规则打分函数
def calculate_rule_rerank_score(item, question):
    score = float(item.get('rrf_score') or item.get('score') or 0)
    if score <= 0:
        score = 0.0001

    text = str(item.get('text') or '')
    title_path = str(item.get('title_path') or '')
    block_type = str(item.get('block_type') or '')
    sources = get_retrieval_source_names(item)
    query_terms = extract_query_terms(question)

    reasons = []

    if {'vector', 'keyword'}.issubset(sources):
        score *= RERANK_DOUBLE_HIT_BONUS
        reasons.append('向量和关键词双命中')

    if title_path and any(term in title_path for term in query_terms):
        score *= RERANK_TITLE_MATCH_BONUS
        reasons.append('标题路径命中问题关键词')

    text_length = len(text)
    if 300 <= text_length <= 1500:
        score *= RERANK_GOOD_LENGTH_BONUS
        reasons.append('片段长度适中')
    elif text_length < 120:
        score *= RERANK_SHORT_TEXT_PENALTY
        reasons.append('片段过短降权')
    elif text_length > 2200:
        score *= RERANK_LONG_TEXT_PENALTY
        reasons.append('片段过长降权')

    if block_type in ['pdf_title', 'title']:
        score *= 0.70
        reasons.append('标题块不能直接回答问题，降权')

    if is_reference_like_chunk(item):
        score *= RERANK_REFERENCE_PENALTY
        reasons.append('参考文献类内容降权')

    if is_explanation_question(question):
        if text.strip().startswith('图') or text.strip().startswith('表'):
            score *= 0.75
            reasons.append('解释型问题图表类内容降权')

    if is_table_or_data_question(question):
        if 'table' in block_type:
            score *= RERANK_TABLE_QUERY_BONUS
            reasons.append('数据型问题优先表格')
    else:
        if 'table' in block_type:
            score *= RERANK_TABLE_MISMATCH_PENALTY
            reasons.append('非数据型问题表格轻微降权')

    reranked_item = {**item, 'rule_rerank_score': score, 'rule_rerank_reasons': reasons}
    return reranked_item

# rerank 列表函数
def rule_rerank_search_results(search_results, question, top_k=RERANK_FINAL_LIMIT):
    reranked_results = [
        calculate_rule_rerank_score(item, question)
        for item in search_results
    ]
    reranked_results.sort(
        key=lambda item: item.get('rule_rerank_score', 0),
        reverse=True,
    )
    return reranked_results[:top_k]

# 使用 rerank 模型
MODEL_RERANK_CANDIDATE_LIMIT = 14
MODEL_RERANK_WEIGHT = 0.60          # 模型rerank权重
RULE_RERANK_WEIGHT = 0.40           # 规则rerank权重

def truncate_rerank_document_text(text, max_chars=3000):
    text = str(text or '').strip()
    if len(text) <= max_chars:
        return text
    return text[:max_chars]

# 规则分数归一化
def normalize_item_scores(items, score_key):
    scores = [
        float(item.get(score_key) or 0)
        for item in items
    ]

    if not scores:
        return {}

    min_score = min(scores)
    max_score = max(scores)
    normalized_scores = {}

    for index, item in enumerate(items):
        score = float(item.get(score_key) or 0)

        if max_score == min_score:
            normalized_scores[index] = 1.0
        else:
            normalized_scores[index] = (score - min_score) / (max_score - min_score)

    return normalized_scores

# 使用模型进行 rerank
def rerank_search_results_with_model(question, search_results, top_k=8):
    if not search_results:
        return []

    api_key = getattr(settings, 'RAG_RERANK_API_KEY', '')
    if not api_key:
        return search_results[:top_k]

    candidate_results = search_results[:MODEL_RERANK_CANDIDATE_LIMIT]
    documents = [
        truncate_rerank_document_text(item.get('text') or '')
        for item in candidate_results
    ]
    payload = {
        'model': getattr(settings, 'RAG_RERANK_MODEL', 'qwen3-rerank'),
        'query': str(question or ''),
        'documents': documents,
        'top_n': min(top_k, len(documents)),
        'instruct': getattr(
            settings,
            'RAG_RERANK_INSTRUCT',
            'Given a web search query, retrieve relevant passages that answer the query.',
        ),
    }

    url = f"{getattr(settings, 'RAG_RERANK_BASE_URL').rstrip('/')}/reranks"

    try:
        response = requests.post(
            url,
            headers={
                'Authorization': f'Bearer {api_key}',
                'Content-Type': 'application/json',
            },
            json=payload,
            timeout=getattr(settings, 'RAG_RERANK_TIMEOUT', 30),
        )
        response.raise_for_status()
        data = response.json()
    except Exception as e:
        print('qwen3-rerank 调用失败:', repr(e))
        return search_results[:top_k]

    results_data = data.get('results') or data.get('output', {}).get('results', [])

    model_score_map = {}

    for result in results_data:
        index = result.get('index')
        if index is None or index >= len(candidate_results):
            continue

        model_score_map[index] = float(result.get('relevance_score') or 0)

    if not model_score_map:
        return search_results[:top_k]

    normalized_rule_scores = normalize_item_scores(
        candidate_results,
        'rule_rerank_score',
    )

    reranked_results = []

    for index, item in enumerate(candidate_results):
        if index not in model_score_map:
            continue

        model_score = model_score_map[index]
        rule_score = normalized_rule_scores.get(index, 0)

        combined_score = (
            model_score * MODEL_RERANK_WEIGHT
            + rule_score * RULE_RERANK_WEIGHT
        )

        reranked_item = {**item, 'model_rerank_score': model_score, 'normalized_rule_rerank_score': rule_score,
                         'model_rule_combined_score': combined_score, 'model_rerank_model': payload['model']}

        reranked_results.append(reranked_item)

    reranked_results.sort(
        key=lambda item: item.get('model_rule_combined_score', 0),
        reverse=True,
    )

    return reranked_results[:top_k]


def qdrant_point_to_search_result(point, score=None):
    payload = point.payload or {}
    metadata = payload.get('metadata') or {}
    text = (
        payload.get('page_content')
        or payload.get('content')
        or payload.get('text')
        or ''
    )

    return {
        'score': score,
        'file_id': metadata.get('file_id'),
        'project_id': metadata.get('project_id'),
        'project_name': metadata.get('project_name'),
        'file_name': metadata.get('file_name'),
        'file_extension': metadata.get('file_extension'),
        'chunk_index': metadata.get('chunk_index'),
        'block_type': metadata.get('block_type'),
        'page': metadata.get('page'),
        'sheet_name': metadata.get('sheet_name'),
        'text': text,
    }
# ———————————————————— 拼prompt ————————————————————
# 抓取相邻 chunk
def fetch_neighbor_chunks_from_qdrant(file_id, chunk_index, neighbor_size=1):
    if file_id is None or chunk_index is None:
        return []

    try:
        chunk_index = int(chunk_index)
    except (TypeError, ValueError):
        return []

    neighbor_indices = list(
        range(
            max(0, chunk_index - neighbor_size),
            chunk_index + neighbor_size + 1,
        )
    )

    client = ensure_langchain_collection()
    qdrant_filter = Filter(
        must=[
            FieldCondition(
                key='metadata.file_id',
                match=MatchValue(value=file_id),
            ),
            FieldCondition(
                key='metadata.chunk_index',
                match=MatchAny(any=neighbor_indices),
            ),
        ]
    )

    points, _ = client.scroll(
        collection_name=settings.LANGCHAIN_QDRANT_COLLECTION,
        scroll_filter=qdrant_filter,
        limit=len(neighbor_indices),
        with_payload=True,
        with_vectors=False,
    )

    results = [qdrant_point_to_search_result(point) for point in points]
    results.sort(key=lambda item: item.get('chunk_index') or 0)
    return results

# 与相邻 chunk 拼接
def expand_search_results_with_neighbors(search_results, neighbor_size=1, allowed_extensions=None):
    expanded_results = []
    seen = set()

    for item in search_results:
        file_extension = item.get('file_extension')

        if allowed_extensions and file_extension not in allowed_extensions:
            neighbor_items = [item]
        else:
            neighbor_items = fetch_neighbor_chunks_from_qdrant(
                file_id=item.get('file_id'),
                chunk_index=item.get('chunk_index'),
                neighbor_size=neighbor_size,
            )

            if not neighbor_items:
                neighbor_items = [item]

        for neighbor in neighbor_items:
            key = (
                neighbor.get('file_id'),
                neighbor.get('chunk_index'),
                neighbor.get('text'),
            )
            if key in seen:
                continue

            seen.add(key)
            expanded_results.append(neighbor)

    return expanded_results

CONTEXT_MAX_CHARS = 9000                # 控制最终塞给大模型的资料总长度
CONTEXT_MAX_ITEMS = 8                   # 最多给模型多少段资料
CONTEXT_MAX_ITEMS_PER_FILE = 5          # 防止同一个文件霸占全部上下文
CONTEXT_MIN_ITEM_CHARS = 80             # 太短的片段一般是标题、页码、残句、图注，将其筛掉
CONTEXT_MAX_ITEM_CHARS = 1800           # 单段资料最大长度

# 安全截断，不让单个片段太长，保证多个片段都有机会进入上下文
def truncate_context_text(text, max_chars=CONTEXT_MAX_ITEM_CHARS):
    text = str(text or '').strip()

    if len(text) <= max_chars:
        return text

    return text[:max_chars].rstrip() + '...'

# 给每个片段生成唯一 key，防重复
def make_context_item_key(item):
    return (
        item.get('file_id'),
        item.get('chunk_index'),
    )

# 按 rerank 结果的优先级给片段取最终排序分数
def get_context_rank_score(item):
    for key in [
        'model_rule_combined_score',
        'model_rerank_score',
        'rule_rerank_score',
        'rrf_score',
        'score',
    ]:
        value = item.get(key)
        if value is not None:
            return float(value)
    return 0

def get_context_display_score(item):
    for key in [
        'model_rule_combined_score',
        'model_rerank_score',
        'normalized_rule_rerank_score',
        'rule_rerank_score',
        'rrf_score',
        'context_rank_score',
        'score',
    ]:
        value = item.get(key)
        if value is not None:
            try:
                return float(value)
            except (TypeError, ValueError):
                continue
    return None

# 低优先级判断函数
def looks_like_table_or_figure_text(text):
    text = str(text or '').strip()
    lines = [line.strip() for line in text.splitlines() if line.strip()]

    if not lines:
        return False

    if text.count('|') >= 4:
        return True

    numeric_lines = sum(1 for line in lines if re.search(r'\d', line))
    short_lines = sum(1 for line in lines if len(line) <= 35)

    if len(lines) >= 5:
        numeric_ratio = numeric_lines / len(lines)
        short_ratio = short_lines / len(lines)

        if numeric_ratio >= 0.45 and short_ratio >= 0.55:
            return True

    figure_keywords = ['图', 'Fig.', 'Fig ', '表', 'Table']
    if any(text.startswith(keyword) for keyword in figure_keywords):
        return True

    return False

def is_low_priority_context_item(item, question):
    if not is_explanation_question(question):
        return False
    block_type = str(item.get('block_type') or '').lower()
    text = str(item.get('text') or '')
    if 'table' in block_type:
        return True
    if 'figure' in block_type or 'caption' in block_type:
        return True
    if looks_like_table_or_figure_text(text):
        return True
    return False

# 去重并排序
def deduplicate_context_items(search_results):
    item_map = {}

    for item in search_results:
        key = make_context_item_key(item)
        if key in [(None, None), None]:
            continue
        current_score = get_context_rank_score(item)
        if key not in item_map:
            item_map[key] = item
            continue
        old_score = get_context_rank_score(item_map[key])
        if current_score > old_score:
            item_map[key] = item

    deduplicated_items = list(item_map.values())
    deduplicated_items.sort(
        key=get_context_rank_score,
        reverse=True,
    )
    return deduplicated_items

# 限制每个文件最多进入多少段
def limit_context_items_per_file(items):
    file_count_map = {}
    limited_items = []

    for item in items:
        file_id = item.get('file_id')
        current_count = file_count_map.get(file_id, 0)

        if current_count >= CONTEXT_MAX_ITEMS_PER_FILE:
            continue

        limited_items.append(item)
        file_count_map[file_id] = current_count + 1

    return limited_items

# 格式化每段资料
def format_context_item(item, index):
    file_name = item.get('file_name') or '未知文件'
    page = item.get('page')
    title_path = item.get('title_path')
    block_type = item.get('block_type')
    score = get_context_display_score(item)
    chunk_indexes = item.get('merged_chunk_indexes') or [item.get('chunk_index')]

    metadata_lines = [
        f'[资料{index}]',
        f'文件：{file_name}',
        f'片段：{", ".join(str(chunk) for chunk in chunk_indexes if chunk is not None)}',
    ]

    if page:
        metadata_lines.append(f'页码：第 {page} 页')
    if title_path:
        metadata_lines.append(f'章节：{title_path}')
    if block_type:
        metadata_lines.append(f'内容类型：{block_type}')
    if score is not None:
        metadata_lines.append(f'相关性分数：{float(score):.4f}')
    text = truncate_context_text(item.get('text'))
    return '\n'.join(metadata_lines) + '\n内容：\n' + text

# 打包构建上下文
def pack_rag_context(search_results, question=None):
    deduplicated_items = deduplicate_context_items(search_results)

    high_priority_items = []
    low_priority_items = []

    for item in deduplicated_items:
        text = str(item.get('text') or '').strip()

        if len(text) < CONTEXT_MIN_ITEM_CHARS:
            continue

        if is_low_priority_context_item(item, question):
            low_priority_items.append(item)
        else:
            high_priority_items.append(item)

    packed_candidates = high_priority_items + low_priority_items
    limited_items = limit_context_items_per_file(packed_candidates)

    packed_parts = []
    selected_items = []
    current_chars = 0

    for item in limited_items:
        formatted_text = format_context_item(
            item,
            index=len(selected_items) + 1,
        )

        if current_chars + len(formatted_text) > CONTEXT_MAX_CHARS:
            continue

        packed_parts.append(formatted_text)
        selected_items.append(item)
        current_chars += len(formatted_text)

        if len(selected_items) >= CONTEXT_MAX_ITEMS:
            break

    context = '\n\n---\n\n'.join(packed_parts)
    return context, selected_items
# ———————————————————— 回答 ————————————————————
# 接入大模型
def get_chat_client():
    if not settings.RAG_CHAT_API_KEY:
        raise ValueError('缺少 RAG_CHAT_API_KEY 配置')

    return OpenAI(
        api_key=settings.RAG_CHAT_API_KEY,
        base_url=settings.RAG_CHAT_BASE_URL,
    )

# 查询不到信息时不显示来源
def should_hide_sources_for_answer(answer):
    no_answer_markers = [
        '资料中没有找到相关信息',
        '没有在项目文件中检索到相关内容',
    ]
    answer_text = answer or ''
    return any(marker in answer_text for marker in no_answer_markers)

# 提取参考资料序号
def extract_cited_source_indexes(answer):
    cited_indexes = re.findall(r'\[资料\s*(\d+)]', answer or '')
    return {
        int(index)
        for index in cited_indexes
        if str(index).isdigit()
    }

# history 清洗，防止用户塞超长文本
def normalize_chat_history(history, max_rounds=5, max_message_chars=300):
    if not isinstance(history, list):
        return []

    valid_messages = []

    for message in history:
        if not isinstance(message, dict):
            continue

        role = message.get('role')
        content = str(message.get('content') or '').strip()

        if role not in ['user', 'assistant']:
            continue

        if not content:
            continue

        valid_messages.append({
            'role': role,
            'content': content[:max_message_chars],
        })

    return valid_messages[-max_rounds * 2:]

# 根据历史问题进行 Query Rewrite
def rewrite_question_with_history(question, history=None):
    history = normalize_chat_history(history)

    if not history:
        return question

    history_text = '\n'.join(
        f"{'用户' if item['role'] == 'user' else '助手'}：{item['content']}"
        for item in history
    )

    prompt = f'''
请根据历史对话，把用户当前问题改写成一个完整、独立、适合项目文档检索的问题。

要求：
1. 只改写问题，不要回答问题。
2. 保留用户当前问题的真实意图。
3. 如果当前问题中有“它”“这个方法”“上面那个”等指代，请根据历史对话补全。
4. 不要添加历史对话中没有依据的新结论。
5. 输出一句改写后的问题。

历史对话：
{history_text}

用户当前问题：
{question}

改写后的独立问题：
'''.strip()

    client = get_chat_client()

    response = client.chat.completions.create(
        model=settings.RAG_CHAT_MODEL,
        messages=[
            {
                'role': 'system',
                'content': '你是一个项目文档问答系统中的查询改写助手，只负责改写检索问题。',
            },
            {
                'role': 'user',
                'content': prompt,
            },
        ],
        temperature=0,
    )

    rewritten = response.choices[0].message.content.strip()

    if not rewritten:
        return question

    return rewritten

# 输出回答
def answer_question_with_rag(question, project_id=None, limit=8, history=None):
    standalone_question = rewrite_question_with_history(question, history)
    search_results = hybrid_search_file_chunks(
        question=standalone_question,
        project_id=project_id,
        final_limit=limit,
    )
    search_results = expand_search_results_with_neighbors(
        search_results,
        neighbor_size=1,
        allowed_extensions={'.pdf'},
    )

    context, packed_results = pack_rag_context(search_results, question=standalone_question)

    if not context:
        yield {
            'type': 'delta',
            'content': '没有在项目文件中检索到相关内容。',
        }
        yield {
            'type': 'done',
            'sources': [],
        }
        return

    prompt = f'''
    你是一个严谨的项目文档问答助手。请只根据下面提供的参考资料回答用户问题。

    回答要求：
    1. 只能根据参考资料回答，不要使用资料外知识。
    2. 每个关键结论后必须标注来源编号，例如：[资料1]、[资料2]。
    3. 如果参考资料无法支持答案，请只回答“资料中没有找到相关信息”。
    4. 不要引用没有实际使用的资料编号。
    5. 回答要清晰、简洁，优先使用中文。

    用户原始问题：
    {question}

    用于检索的独立问题：
    {standalone_question}

    参考资料：
    {context}
    '''.strip()

    client = get_chat_client()

    stream = client.chat.completions.create(
        model=settings.RAG_CHAT_MODEL,
        messages=[
            {
                'role': 'system',
                'content': '你是一个严谨的项目文档问答助手，只能基于给定资料回答。',
            },
            {
                'role': 'user',
                'content': prompt,
            },
        ],
        temperature=0.2,
        stream=True,
    )

    answer_parts = []

    for chunk in stream:
        choices = getattr(chunk, 'choices', None) or []
        if not choices:
            continue

        delta = getattr(choices[0], 'delta', None)
        content = getattr(delta, 'content', None)
        if content:
            answer_parts.append(content)
            yield {
                'type': 'delta',
                'content': content,
            }

    if should_hide_sources_for_answer(''.join(answer_parts)):
        packed_results = []

    answer_text = ''.join(answer_parts)

    if should_hide_sources_for_answer(answer_text):
        packed_results = []

    cited_indexes = extract_cited_source_indexes(answer_text)

    sources = []
    seen = set()

    for source_index, item in enumerate(packed_results, start=1):
        if source_index not in cited_indexes:
            continue

        key = (
            item.get('file_id'),
            tuple(item.get('merged_chunk_indexes') or [item.get('chunk_index')]),
        )
        if key in seen:
            continue
        seen.add(key)

        sources.append({
            'source_index': source_index,
            'file_id': item.get('file_id'),
            'file_name': item.get('file_name'),
            'score': get_context_display_score(item),
            'chunk_index': item.get('chunk_index'),
            'merged_chunk_indexes': item.get('merged_chunk_indexes'),
            'page': item.get('page'),
            'sheet_name': item.get('sheet_name'),
        })
    yield {
        'type': 'done',
        'sources': sources,
    }
