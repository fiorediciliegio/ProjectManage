from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import load_workbook

from app01.services.rag.blocks import make_block


# 纯文本解析
def load_plain_file_blocks(file_path, base_metadata):
    text = Path(file_path).read_text(encoding='utf-8',errors='replace').strip()

    if not text:
        return []

    return [
        make_block(
            text=text,
            block_type='plain_text',
            base_metadata=base_metadata,
            order=0,
        )
    ]

# .docx 解析
def get_heading_level(paragraph):
    style_name = paragraph.style.name if paragraph.style else ''
    heading_map = {
        'Heading 1': 1,
        'Heading 2': 2,
        'Heading 3': 3,
        '标题 1': 1,
        '标题 2': 2,
        '标题 3': 3,
    }
    return heading_map.get(style_name)

def load_docx_blocks(file_path, base_metadata):
    document = DocxDocument(file_path)
    blocks = []
    title_stack = []
    order = 0

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        heading_level = get_heading_level(paragraph)
        if heading_level:
            title_stack = title_stack[:heading_level - 1]
            title_stack.append(text)
            blocks.append(
                make_block(
                    text=text,
                    block_type='title',
                    base_metadata=base_metadata,
                    order=order,
                    title_path=title_stack.copy(),
                )
            )
        else:
            blocks.append(
                make_block(
                    text=text,
                    block_type='paragraph',
                    base_metadata=base_metadata,
                    order=order,
                    title_path=title_stack.copy(),
                )
            )
        order += 1
    for table_index, table in enumerate(document.tables):
        table_text = convert_docx_table_to_markdown(table)
        if not table_text.strip():
            continue
        blocks.append(
            make_block(
                text=table_text,
                block_type='table',
                base_metadata=base_metadata,
                order=order,
                title_path=title_stack.copy(),
                caption=f'Word 表格 {table_index + 1}',
            )
        )
        order += 1
    return blocks

# .docx 表格转 Markdown
def convert_docx_table_to_markdown(table):
    rows = []

    for row in table.rows:
        cells = [
            cell.text.strip().replace('\n', ' ')
            for cell in row.cells
        ]
        if any(cells):
            rows.append(cells)

    if not rows:
        return ''

    max_cols = max(len(row) for row in rows)
    normalized_rows = [
        row + [''] * (max_cols - len(row))
        for row in rows
    ]

    header = normalized_rows[0]
    body = normalized_rows[1:]

    lines = []
    lines.append('| ' + ' | '.join(header) + ' |')
    lines.append('| ' + ' | '.join(['---'] * max_cols) + ' |')

    for row in body:
        lines.append('| ' + ' | '.join(row) + ' |')

    return '\n'.join(lines)

# .xlsx 解析
def load_xlsx_blocks(file_path, base_metadata):
    value_workbook = load_workbook(file_path, data_only=True)
    formula_workbook = load_workbook(file_path, data_only=False)
    blocks = []
    order = 0
    for sheet_name in value_workbook.sheetnames:
        value_sheet = value_workbook[sheet_name]
        formula_sheet = formula_workbook[sheet_name]

        table_text = convert_xlsx_sheet_to_markdown_with_formulas(
            value_sheet=value_sheet,
            formula_sheet=formula_sheet,
        )
        if not table_text.strip():
            continue
        blocks.append(
            make_block(
                text=table_text,
                block_type='sheet_table',
                base_metadata=base_metadata,
                order=order,
                sheet_name=sheet_name,
                caption=f'工作表：{sheet_name}',
            )
        )
        order += 1
    return blocks

# 表格转换
def convert_xlsx_sheet_to_markdown_with_formulas(value_sheet, formula_sheet):
    rows = []
    max_row = max(value_sheet.max_row, formula_sheet.max_row)
    max_column = max(value_sheet.max_column, formula_sheet.max_column)

    for row_index in range(1, max_row + 1):
        row_cells = []

        for column_index in range(1, max_column + 1):
            value_cell = value_sheet.cell(row=row_index, column=column_index)
            formula_cell = formula_sheet.cell(row=row_index, column=column_index)

            display_text = format_excel_cell_value_with_formula(
                value=formula_cell.value,
                calculated_value=value_cell.value,
            )

            row_cells.append(display_text)

        if any(cell.strip() for cell in row_cells):
            rows.append(row_cells)

    if not rows:
        return ''

    max_cols = max(len(row) for row in rows)

    normalized_rows = [
        row + [''] * (max_cols - len(row))
        for row in rows
    ]

    header = normalized_rows[0]
    body = normalized_rows[1:]

    lines = []
    lines.append('| ' + ' | '.join(header) + ' |')
    lines.append('| ' + ' | '.join(['---'] * max_cols) + ' |')

    for row in body:
        lines.append('| ' + ' | '.join(row) + ' |')
    return '\n'.join(lines)

# 单元格格式化
def format_excel_cell_value_with_formula(value, calculated_value):
    if value is None and calculated_value is None:
        return ''

    value_text = '' if value is None else str(value).strip()
    calculated_text = '' if calculated_value is None else str(calculated_value).strip()

    if value_text.startswith('='):
        if calculated_text:
            return f'{value_text}（计算结果：{calculated_text}）'
        return value_text

    return value_text
