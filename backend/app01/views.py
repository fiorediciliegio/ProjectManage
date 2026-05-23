from django.views.decorators.csrf import ensure_csrf_cookie
from django.contrib.auth import authenticate, login, logout
from rest_framework.decorators import api_view
from rest_framework import status
from django.http import JsonResponse
from django.http import HttpResponse
from django.http import StreamingHttpResponse
from django.utils.encoding import smart_str
from .models import AuditLog
from .models import Project
from .models import ProjectNode
from .models import Person
from .models import File
from .models import QualityInspectionTemplate
from .models import QualityInspectionReport
from .models import InspectionItem
from .models import InspectionResult
from .models import SecurityCheckTemplate
from .models import SecurityCheckItem
from .models import SecurityInspectionReport
from .models import SecurityCheckResult
from .models import SecurityInspectionImage
from .models import SafetyIssueSolution
from .models import CostInformation
from .serializers import (
    ProjectSerializer,
    PersonSerializer,
    FileSerializer,
    CostInformationSerializer,
    QualityInspectionTemplateSerializer,
    QualityInspectionReportSerializer,
    SecurityCheckTemplateSerializer,
    SecurityInspectionReportSerializer,
    SafetyIssueSolutionSerializer,
)
from .response import success_response, error_response
from django.db.models.functions import ExtractQuarter
import os
import json
import datetime
from django.db.models import Count
import mimetypes
from docx import Document
from django.db.models import Sum
from django.db.models.functions import TruncMonth
from html import escape
from openpyxl import load_workbook
from app01.services.elasticsearch_service import search_audit_logs, delete_file_chunks_from_elasticsearch
from app01.services.langchain_rag_service import (
    index_file_to_qdrant_langchain,
    get_indexed_file_ids,
    answer_question_with_rag,
    delete_langchain_file_vectors
)

# ———————————————————— 登录 ————————————————————
# 登录
@api_view(['POST'])
def login_view(request):
    try:
        username = request.data.get('username')
        password = request.data.get('password')

        if not username or not password:
            return error_response(
                message='用户名和密码不能为空',
                status_code=status.HTTP_400_BAD_REQUEST,
            )
        user = authenticate(request, username=username, password=password)
        if user is None:
            return error_response(
                message='用户名或密码错误',
                status_code=status.HTTP_400_BAD_REQUEST,
            )
        if not user.is_active:
            return error_response(
                message='账号已被禁用',
                status_code=status.HTTP_403_FORBIDDEN,
            )
        login(request, user)
        create_audit_log(
            request=request,
            action='login',
            module='用户',
            target_id=user.pk,
            target_name=username,
            description='登录'
        )
        try:
            person = user.person_profile
            person_data = {
                'person_id': person.id,
                'person_name': person.NAME_Person,
                'person_role': person.POS_Person,
                'sys_role': person.sys_role,
            }
        except Person.DoesNotExist:
            person_data = {
                'person_id': None,
                'person_name': None,
                'person_role': None,
                'sys_role': None,
            }
        return success_response(
            data={
                'username': user.username,
                'is_superuser': user.is_superuser,
                'is_staff': user.is_staff,
                **person_data,
            },
            message='登录成功',
            status_code=status.HTTP_200_OK,
        )
    except Exception as e:
        return error_response(
            message=str(e),
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )

# 获取当前登录用户信息
@api_view(['GET'])
def current_user_view(request):
    if not request.user.is_authenticated:
        return error_response(
            message='当前未登录',
            status_code=status.HTTP_401_UNAUTHORIZED,
        )
    user = request.user
    try:
        person = user.person_profile
        person_data = {
            'person_id': person.id,
            'person_name': person.NAME_Person,
            'person_role': person.POS_Person,
            'sys_role': person.sys_role,
        }
    except Person.DoesNotExist:
        person_data = {
            'person_id': None,
            'person_name': None,
            'person_role': None,
            'sys_role': None,
        }
    return success_response(
        data={
            'username': user.username,
            'is_superuser': user.is_superuser,
            'is_staff': user.is_staff,
            **person_data,
        },
        message='获取当前用户信息成功',
        status_code=status.HTTP_200_OK,
    )

# 退出登录
@api_view(['POST'])
def logout_view(request):
    user = request.user
    create_audit_log(
        request=request,
        action='logout',
        module='用户',
        target_id=user.id,
        target_name=user.username,
        description='退出登录',
    )
    logout(request)
    return success_response(
        data=None,
        message='退出登录成功',
        status_code=status.HTTP_200_OK,
    )

# 检验登录状态
def get_authenticated_person(request):
    if not request.user.is_authenticated:
        return None, error_response(
            message='请先登录',
            status_code=status.HTTP_401_UNAUTHORIZED,
        )
    try:
        person = request.user.person_profile
        return person, None
    except Person.DoesNotExist:
        return None, error_response(
            message='当前账号未绑定人员信息',
            status_code=status.HTTP_403_FORBIDDEN,
        )

# 检验账号权限
def is_admin(person):
    return person.sys_role == 'admin'

def is_project_manager(person):
    return person.sys_role == 'project_manager'

def is_admin_or_project_manager(person):
    return person.sys_role in ['admin', 'project_manager']

def permission_denied(message='无权限操作'):
    return error_response(
        message=message,
        status_code=status.HTTP_403_FORBIDDEN,
    )

def is_project_member(project, person):
    return project.ID_Person.filter(id=person.id).exists()

def is_project_manager_of(project, person):
    return (
        person.sys_role == 'project_manager'
        and person.POS_Person == '项目经理'
        and is_project_member(project, person)
    )

def is_project_budgeter_of(project, person):
    return (
        person.POS_Person in ['预算员', '商务经理']
        and is_project_member(project, person)
    )

def is_project_quality_officer_of(project, person):
    return (
        person.POS_Person == '质量员'
        and is_project_member(project, person)
    )

def is_project_safety_officer_of(project, person):
    return (
        person.POS_Person in ['安全员', '安全经理']
        and is_project_member(project, person)
    )

def is_project_file_manager_of(project, person):
    return (
        person.POS_Person in ['资料员', '资料主管']
        and is_project_member(project, person)
    )

def can_manage_project(project, person):
    return is_admin(person) or is_project_manager_of(project, person)

def can_manage_project_cost(project, person):
    return (
        is_admin(person)
        or is_project_manager_of(project, person)
        or is_project_budgeter_of(project, person)
    )

def can_manage_project_quality(project, person):
    return (
        is_admin(person)
        or is_project_manager_of(project, person)
        or is_project_quality_officer_of(project, person)
    )

def can_manage_project_safety(project, person):
    return (
        is_admin(person)
        or is_project_manager_of(project, person)
        or is_project_safety_officer_of(project, person)
    )

def can_manage_project_file(project, person):
    return (
        is_admin(person)
        or is_project_manager_of(project, person)
        or is_project_file_manager_of(project, person)
    )

@ensure_csrf_cookie
@api_view(['GET'])
def csrf_token_view(request):
    return success_response(
        data=None,
        message='CSRF cookie 设置成功',
        status_code=status.HTTP_200_OK,
    )

# ———————————————————— 日志 ————————————————————
# 创建日志
def create_audit_log(request, action, module, target_id=None, target_name='', description=''):
    user = request.user if request.user.is_authenticated else None
    person = getattr(user, 'person_profile', None) if user else None
    log = AuditLog.objects.create(
        user=user,
        person=person,
        action=action,
        module=module,
        target_id=target_id,
        target_name=target_name,
        description=description,
    )
    try:
        from app01.services.elasticsearch_service import index_audit_log
        index_audit_log(log)
    except Exception as e:
        print(f'操作日志写入 Elasticsearch 失败：{e}')
    return log

# ES 查看并操作日志
@api_view(['GET'])
def rest_audit_logs_search(request):
    person, auth_error = get_authenticated_person(request)
    if auth_error:
        return auth_error

    if not is_admin(person):
        return permission_denied('只有管理员可以查看操作日志')

    keyword = request.GET.get('keyword', '').strip()
    module = request.GET.get('module', '').strip()
    action = request.GET.get('action', '').strip()
    date = request.GET.get('date', '').strip()

    try:
        logs = search_audit_logs(
            keyword=keyword,
            module=module,
            action=action,
            date=date,
            size=100,
        )
        return success_response(
            data={'logs': logs},
            message='获取操作日志成功',
            status_code=status.HTTP_200_OK,
        )
    except Exception as e:
        return error_response(
            message=str(e),
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )

# ———————————————————— 项目与人员 ————————————————————
# REST 项目模块：项目列表、详情、新增与删除接口
@api_view(['GET', 'POST', 'PATCH', 'DELETE'])
def rest_projects(request, project_id=None):
    if request.method == 'GET':
        try:
            if project_id is not None:
                project = Project.objects.get(id=project_id)
                serializer = ProjectSerializer(project)
                return success_response(
                    data={'project': serializer.data},
                    message='获取项目详情成功',
                    status_code=status.HTTP_200_OK,
                )
            projects = Project.objects.all()
            serializer = ProjectSerializer(projects, many=True)
            return success_response(
                data={'projects': serializer.data},
                message='获取项目列表成功',
                status_code=status.HTTP_200_OK,
            )
        except Project.DoesNotExist:
            return error_response(
                message='项目不存在',
                status_code=status.HTTP_404_NOT_FOUND,
            )
        except Exception as e:
            return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

    if request.method == 'POST':
        # 登录校验
        person, auth_error = get_authenticated_person(request)
        if auth_error:
            return auth_error
        if not is_admin(person):
            return permission_denied('只有管理员可以创建项目')

        if project_id is not None:
            return error_response(
                message='创建项目请使用项目集合接口',
                status_code=status.HTTP_405_METHOD_NOT_ALLOWED,
            )
        project_manager_id = request.data.get('pjmanager_id')
        if not project_manager_id:
            return error_response(
                message='请选择项目负责人',
                data={'pjmanager_id': '请选择项目负责人'},
                status_code=status.HTTP_400_BAD_REQUEST,
            )
        try:
            project_manager = Person.objects.get(id=project_manager_id, POS_Person='项目经理')
        except Person.DoesNotExist:
            return error_response(
                message='项目负责人必须从项目经理中选择',
                data={'pjmanager_id': '项目负责人必须从项目经理中选择'},
                status_code=status.HTTP_400_BAD_REQUEST,
            )

        project_data = request.data.copy()
        project_data['pjmanager'] = project_manager.NAME_Person
        project_data.pop('pjmanager_id', None)

        serializer = ProjectSerializer(data=project_data)
        if serializer.is_valid():
            project = serializer.save()
            # 创建日志
            create_audit_log(
                request=request,
                action='create',
                module='项目管理',
                target_id=project.id,
                target_name=project.NAME_Project,
                description=f'创建项目：{project.NAME_Project}',
            )
            project.ID_Person.add(project_manager)
            return success_response(
                data={'project': ProjectSerializer(project).data},
                message='项目创建成功',
                status_code=status.HTTP_201_CREATED,
            )
        return error_response(
            message='项目创建失败',
            data=serializer.errors,
            status_code=status.HTTP_400_BAD_REQUEST,
        )
    if request.method == 'PATCH':
        if project_id is None:
            return error_response(
                message='缺少项目ID',
                status_code=status.HTTP_400_BAD_REQUEST,
            )
        try:
            project = Project.objects.get(id=project_id)
        except Project.DoesNotExist:
            return error_response(
                message='项目不存在',
                status_code=status.HTTP_404_NOT_FOUND,
            )
        # 登录校验
        person, auth_error = get_authenticated_person(request)
        if auth_error:
            return auth_error
        if not can_manage_project(project, person):
            return permission_denied('只有管理员或当前项目经理可以修改项目详情')

        project_data = request.data.copy()
        project_manager_id = project_data.get('pjmanager_id')
        if project_manager_id:
            try:
                project_manager = Person.objects.get(id=project_manager_id, POS_Person='项目经理')
            except Person.DoesNotExist:
                return error_response(
                    message='项目负责人必须从项目经理中选择',
                    data={'pjmanager_id': '项目负责人必须从项目经理中选择'},
                    status_code=status.HTTP_400_BAD_REQUEST,
                )
            project_data['pjmanager'] = project_manager.NAME_Person
            project_data.pop('pjmanager_id', None)
        else:
            project_manager = None

        serializer = ProjectSerializer(project, data=project_data, partial=True)
        if serializer.is_valid():
            project = serializer.save()
            # 创建日志
            create_audit_log(
                request=request,
                action='update',
                module='项目管理',
                target_id=project.id,
                target_name=project.NAME_Project,
                description=f'修改项目详情：{project.NAME_Project}',
            )
            if project_manager is not None:
                project.ID_Person.add(project_manager)
            return success_response(
                data={'project': ProjectSerializer(project).data},
                message='项目详情修改成功',
                status_code=status.HTTP_200_OK,
            )
        return error_response(
            message='项目详情修改失败',
            data=serializer.errors,
            status_code=status.HTTP_400_BAD_REQUEST,
        )
    # 登录校验
    person, auth_error = get_authenticated_person(request)
    if auth_error:
        return auth_error
    if not is_admin(person):
        return permission_denied('只有管理员可以删除项目')

    if project_id is None:
        return error_response(
            message='缺少项目ID',
            status_code=status.HTTP_400_BAD_REQUEST,
        )
    try:
        project = Project.objects.get(id=project_id)
        project_name = project.NAME_Project
        project.delete()
        create_audit_log(
            request=request,
            action='delete',
            module='项目管理',
            target_id=project.id,
            target_name=project_name,
            description=f'删除项目：{project_name}',
        )
        return success_response(
            data=None,
            message='项目删除成功',
            status_code=status.HTTP_200_OK,
        )
    except Project.DoesNotExist:
        return error_response(
            message='项目不存在',
            status_code=status.HTTP_404_NOT_FOUND,
        )
    except Exception as e:
        return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

# 序列化项目节点，统一返回给前端的字段结构
def _serialize_project_node(project_node):
    return {
        'pjn_id': project_node.id,
        'pjn_name': project_node.NAME_Milestone,
        'pjn_ddl': str(project_node.DDL_Milestone),
        'pjn_des': project_node.DES_Milestone,
        'pjn_status': project_node.PHEN_Milestone,
    }

# 统计项目节点状态
@api_view(['GET'])
def projectnode_collect(request, project_id=None):
    try:
        if project_id is None:
            return error_response(
                message='缺少项目ID',
                status_code=status.HTTP_400_BAD_REQUEST,
            )
        summary = {
            'completed_count': ProjectNode.objects.filter(ID_Project=project_id, PHEN_Milestone='已完成').count(),
            'in_progress_count': ProjectNode.objects.filter(ID_Project=project_id, PHEN_Milestone='进行中').count(),
            'pending_count': ProjectNode.objects.filter(ID_Project=project_id, PHEN_Milestone='未处理').count(),
        }
        return success_response(
            data={'summary': summary},
            message='获取项目节点统计成功',
            status_code=status.HTTP_200_OK,
        )
    except Exception as e:
        return error_response(
            message=str(e),
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )

# 判断某个项目中是否存在指定职位的人员
def _project_person_exists(project_id, person_name, role):
    if not person_name:
        return False
    return Project.objects.filter(
        id=project_id,
        ID_Person__NAME_Person=person_name,
        ID_Person__POS_Person=role,
    ).exists()

# 统计各职位人员数量
@api_view(['GET'])
def person_collect(request, project_id=None):
    try:
        positions = [
            '项目经理', '生产经理', '技术总工', '安全经理', '商务经理',
            '材料主管', '资料主管', '综合办主任', '工程师', '技术员',
            '质量员', '预算员', '安全员', '资料员', '施工员'
        ]
        count_data = {position: 0 for position in positions}

        if project_id:
            project = Project.objects.get(id=project_id)
            persons = project.ID_Person.filter(POS_Person__in=positions)
        else:
            persons = Person.objects.filter(POS_Person__in=positions)

        for person in persons:
            count_data[person.POS_Person] += 1

        return success_response(
            data={'counts': count_data},
            message='获取人员统计成功',
            status_code=status.HTTP_200_OK,
        )
    except Exception as e:
        return error_response(
            message=str(e),
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )

# REST 项目节点模块：按项目获取节点列表
@api_view(['GET'])
def rest_project_nodes_by_project(request, project_id):
    try:
        projectnodes = ProjectNode.objects.filter(ID_Project=project_id)
        projectnodes_data = [_serialize_project_node(projectnode) for projectnode in projectnodes]
        return success_response(
            data={'project_nodes': projectnodes_data},
            message='获取项目节点列表成功',
            status_code=status.HTTP_200_OK,
        )
    except Exception as e:
        return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

# REST 项目节点模块：新增、删除和更新节点状态接口
@api_view(['POST', 'DELETE', 'PATCH', 'PUT'])
def rest_project_nodes(request, node_id=None):
    if request.method == 'POST':
        try:
            name_projectnode = request.data.get('pjn_name')
            ddl_projectnode_str = request.data.get('pjn_ddl')
            des_projectnode = request.data.get('pjn_des')
            phen_projectnode = request.data.get('pjn_status')
            id_project = request.data.get('pj_id')
            if not ddl_projectnode_str:
                return error_response(
                    message='节点日期不能为空',
                    status_code=status.HTTP_400_BAD_REQUEST,
                )
            ddl_projectnode = datetime.datetime.strptime(ddl_projectnode_str, '%Y-%m-%d').date()
            project_instance = Project.objects.get(id=id_project)
            # 登录校验
            person, auth_error = get_authenticated_person(request)
            if auth_error:
                return auth_error
            # 权限校验
            if not can_manage_project(project_instance, person):
                return permission_denied('只有管理员或当前项目经理可以操作')

            project_node = ProjectNode.objects.create(
                NAME_Milestone=name_projectnode,
                DDL_Milestone=ddl_projectnode,
                DES_Milestone=des_projectnode,
                PHEN_Milestone=phen_projectnode,
                ID_Project=project_instance,
            )
            # 创建日志
            create_audit_log(
                request=request,
                action='create',
                module='项目节点',
                target_id=project_instance.id,
                target_name=project_instance.NAME_Project,
                description=f'创建项目节点：{name_projectnode}',
            )
            return success_response(
                data={'project_node': _serialize_project_node(project_node)},
                message='项目节点创建成功',
                status_code=status.HTTP_201_CREATED,
            )
        except Project.DoesNotExist:
            return error_response(
                message='项目不存在',
                status_code=status.HTTP_404_NOT_FOUND,
            )
        except Exception as e:
            return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

    if node_id is None:
        return error_response(
            message='缺少项目节点ID',
            status_code=status.HTTP_400_BAD_REQUEST,
        )
    try:
        project_node = ProjectNode.objects.get(id=node_id)
    except ProjectNode.DoesNotExist:
        return error_response(
            message='项目节点不存在',
            status_code=status.HTTP_404_NOT_FOUND,
        )

    if request.method == 'DELETE':
        person, auth_error = get_authenticated_person(request)
        if auth_error:
            return auth_error
        # 权限校验
        project = project_node.ID_Project
        if not can_manage_project(project, person):
            return permission_denied('只有管理员或当前项目经理可以操作')
        create_audit_log(
            request=request,
            action='delete',
            module='项目节点',
            target_id=project.id,
            target_name=project.NAME_Project,
            description=f'删除节点：{project_node.NAME_Milestone}',
        )
        project_node.delete()
        return success_response(
            data=None,
            message='项目节点删除成功',
            status_code=status.HTTP_200_OK,
        )
    # 登录校验
    person, auth_error = get_authenticated_person(request)
    if auth_error:
        return auth_error
    # 权限校验
    project = project_node.ID_Project
    if not can_manage_project(project, person):
        return permission_denied('只有管理员或当前项目经理可以操作')

    new_status = request.data.get('pjn_status') or request.data.get('new_pjn_status')
    if not new_status:
        return error_response(
            message='节点状态不能为空',
            status_code=status.HTTP_400_BAD_REQUEST,
        )
    project_node.PHEN_Milestone = new_status
    project_node.save()
    create_audit_log(
        request=request,
        action='update',
        module='项目节点',
        target_id=project.id,
        target_name=project.NAME_Project,
        description=f'节点状态更新：{project_node.NAME_Milestone}',
    )
    return success_response(
        data={'project_node': _serialize_project_node(project_node)},
        message='项目节点状态更新成功',
        status_code=status.HTTP_200_OK,
    )

# REST 人员模块：人员列表、详情、新增与删除接口
@api_view(['GET', 'POST', 'DELETE'])
def rest_persons(request, person_id=None):
    if request.method == 'GET':
        try:
            if person_id is not None:
                person = Person.objects.get(id=person_id)
                serializer = PersonSerializer(person)
                return success_response(
                    data={'person': serializer.data},
                    message='获取人员详情成功',
                    status_code=status.HTTP_200_OK,
                )
            persons = Person.objects.all()
            serializer = PersonSerializer(persons, many=True)
            return success_response(
                data={'persons': serializer.data},
                message='获取人员列表成功',
                status_code=status.HTTP_200_OK,
            )
        except Person.DoesNotExist:
            return error_response(
                message='人员不存在',
                status_code=status.HTTP_404_NOT_FOUND,
            )
        except Exception as e:
            return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

    if request.method == 'POST':
        # 登录校验
        person, auth_error = get_authenticated_person(request)
        if auth_error:
            return auth_error
        if not is_admin(person):
            return permission_denied('只有管理员可以添加人员')
        if person_id is not None:
            return error_response(
                message='创建人员请使用人员集合接口',
                status_code=status.HTTP_405_METHOD_NOT_ALLOWED,
            )
        serializer = PersonSerializer(data=request.data)
        if serializer.is_valid():
            person = serializer.save()
            create_audit_log(
                request=request,
                action='create',
                module='人员管理',
                target_id=person.id,
                target_name=person.NAME_Person,
                description=f'新增人员：{person.NAME_Person}',
            )
            return success_response(
                data={'person': PersonSerializer(person).data},
                message='人员创建成功',
                status_code=status.HTTP_201_CREATED,
            )
        return error_response(
            message='人员创建失败',
            data=serializer.errors,
            status_code=status.HTTP_400_BAD_REQUEST,
        )
    # 登录校验
    person, auth_error = get_authenticated_person(request)
    if auth_error:
        return auth_error
    if not is_admin(person):
        return permission_denied('只有管理员可以删除人员')

    if person_id is None:
        return error_response(
            message='缺少人员ID',
            status_code=status.HTTP_400_BAD_REQUEST,
        )
    try:
        person = Person.objects.get(id=person_id)
        create_audit_log(
            request=request,
            action='delete',
            module='人员管理',
            target_id=person.id,
            target_name=person.NAME_Person,
            description=f'删除人员：{person.NAME_Person}',
        )
        person.delete()
        return success_response(
            data=None,
            message='人员删除成功',
            status_code=status.HTTP_200_OK,
        )
    except Person.DoesNotExist:
        return error_response(
            message='人员不存在',
            status_code=status.HTTP_404_NOT_FOUND,
        )
    except Exception as e:
        return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

# REST 项目人员模块：项目人员列表与添加接口
@api_view(['GET', 'POST'])
def rest_project_persons(request, project_id):
    if request.method == 'GET':
        try:
            project = Project.objects.get(id=project_id)
            persons = project.ID_Person.all()
            serializer = PersonSerializer(persons, many=True)
            return success_response(
                data={'persons': serializer.data},
                message='获取项目人员列表成功',
                status_code=status.HTTP_200_OK,
            )
        except Project.DoesNotExist:
            return error_response(message='项目不存在', status_code=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)
    # 登录校验
    person, auth_error = get_authenticated_person(request)
    if auth_error:
        return auth_error
    # 权限校验
    project = Project.objects.get(id=project_id)
    if not can_manage_project(project, person):
        return permission_denied('只有管理员或当前项目经理可以操作')

    person_ids = request.data.get('person_ids', None)
    if project_id is None or person_ids is None:
        return error_response(
            message='缺少项目ID或人员ID',
            status_code=status.HTTP_400_BAD_REQUEST,
        )
    persons = Person.objects.filter(id__in=person_ids)
    project.ID_Person.add(*persons)
    # 日志
    person_names = '、'.join(f'{person.NAME_Person}(编号:{person.NUM_Person})'for person in persons)
    create_audit_log(
        request=request,
        action='create',
        module='人员管理',
        target_id=project.id,
        target_name=project.NAME_Project,
        description=f'向项目内添加人员：{person_names}',
    )
    return success_response(
        data=None,
        message='项目人员添加成功',
        status_code=status.HTTP_200_OK,
    )

# REST 项目人员模块：移除项目人员接口
@api_view(['DELETE'])
def rest_project_person_detail(request, project_id, person_id):
    # 登录校验
    person, auth_error = get_authenticated_person(request)
    if auth_error:
        return auth_error
    # 权限校验
    project = Project.objects.get(id=project_id)
    if not can_manage_project(project, person):
        return permission_denied('只有管理员或当前项目经理可以操作')
    try:
        remove_person = Person.objects.get(id=person_id)
    except Person.DoesNotExist:
        return error_response(
            message='项目人员详情不存在',
            status_code=status.HTTP_404_NOT_FOUND,
        )
    if not project.ID_Person.filter(id=person_id).exists():
        return error_response(
            message='项目人员关系不存在',
            status_code=status.HTTP_404_NOT_FOUND,
        )
    project.ID_Person.remove(remove_person)
    create_audit_log(
        request=request,
        action='delete',
        module='人员管理',
        target_id=project.id,
        target_name=project.NAME_Project,
        description=f'从项目内移除人员：{remove_person.NAME_Person}(编号:{remove_person.NUM_Person})',
    )
    return success_response(
        data=None,
        message='项目人员移除成功',
        status_code=status.HTTP_200_OK,
    )

# ———————————————————— 成本控制 ————————————————————
# 统计预算与执行成本
@api_view(['GET'])
def project_cost_summary(request, project_id):
    try:
        Project.objects.get(id=project_id)
        cost_items = CostInformation.objects.filter(ID_Project_id=project_id)
        expense_types = ['材料费用', '设备费用', '人工费用', '管理费用', '规费税金', '其他费用']

        result = {
            'TotalBudget': 0,
            'TotalCost': 0,
            'detail': {},
        }
        for expense_type in expense_types:
            result['detail'][expense_type] = {
                'totalbudget': 0,
                'totalcost': 0,
            }

        for expense_type in expense_types:
            budget_summary = cost_items.filter(TYPE_Expense=expense_type).aggregate(
                total_budget_amount=Sum('BUDGET_Amount')
            )
            cost_summary = cost_items.filter(TYPE_Expense=expense_type).aggregate(
                total_cost_amount=Sum('COST_Amount')
            )

            result['detail'][expense_type]['totalbudget'] = budget_summary['total_budget_amount'] or 0
            result['detail'][expense_type]['totalcost'] = cost_summary['total_cost_amount'] or 0
            result['TotalBudget'] += result['detail'][expense_type]['totalbudget']
            result['TotalCost'] += result['detail'][expense_type]['totalcost']

        return success_response(
            data={'summary': result},
            message='获取成本汇总成功',
            status_code=status.HTTP_200_OK,
        )
    except Project.DoesNotExist:
        return error_response(
            message='项目不存在',
            status_code=status.HTTP_404_NOT_FOUND,
        )
    except Exception as e:
        return error_response(
            message=str(e),
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )

# 按月统计实际成本
@api_view(['GET'])
def project_monthly_cost_summary(request, project_id):
    try:
        Project.objects.get(id=project_id)
        cost_items = CostInformation.objects.filter(ID_Project_id=project_id)
        expense_types = ['材料费用', '设备费用', '人工费用', '管理费用', '规费税金', '其他费用']
        monthly_costs = {}

        for expense_type in expense_types:
            monthly_data = (
                cost_items.filter(TYPE_Expense=expense_type)
                .annotate(month=TruncMonth('DATE_Cost'))
                .values('month')
                .annotate(total_cost_amount=Sum('COST_Amount'))
                .values('month', 'total_cost_amount')
            )

            for item in monthly_data:
                month = item['month'].strftime('%Y-%m')
                if month not in monthly_costs:
                    monthly_costs[month] = {et: 0 for et in expense_types}
                monthly_costs[month][expense_type] = item['total_cost_amount'] or 0

        result = [
            {
                'month': month,
                'material': costs['材料费用'],
                'equipment': costs['设备费用'],
                'labour': costs['人工费用'],
                'manage': costs['管理费用'],
                'tax': costs['规费税金'],
                'other': costs['其他费用'],
            }
            for month, costs in monthly_costs.items()
        ]

        return success_response(
            data={'monthlyCosts': result},
            message='获取月度成本统计成功',
            status_code=status.HTTP_200_OK,
        )
    except Project.DoesNotExist:
        return error_response(
            message='项目不存在',
            status_code=status.HTTP_404_NOT_FOUND,
        )
    except Exception as e:
        return error_response(
            message=str(e),
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )

# REST 成本模块：项目成本单列表与新增接口
@api_view(['GET', 'POST'])
def rest_project_costs(request, project_id):
    if request.method == 'GET':
        try:
            project = Project.objects.get(id=project_id)
            costs = CostInformation.objects.filter(ID_Project=project)
            serializer = CostInformationSerializer(costs, many=True)
            return success_response(
                data={'costs': serializer.data},
                message='获取成本单列表成功',
                status_code=status.HTTP_200_OK,
            )
        except Project.DoesNotExist:
            return error_response(
                message='项目不存在',
                status_code=status.HTTP_404_NOT_FOUND,
            )
        except Exception as e:
            return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)
    try:
        project = Project.objects.get(id=project_id)
    except Project.DoesNotExist:
        return error_response(
            message='项目不存在',
            status_code=status.HTTP_404_NOT_FOUND,
        )
    # 登录校验
    person, auth_error = get_authenticated_person(request)
    if auth_error:
        return auth_error
    # 权限校验
    if not can_manage_project_cost(project, person):
        return permission_denied('只有管理员、当前项目经理或当前项目财务人员可以操作成本单')

    data = request.data
    expense_type = data.get('expenseType')
    accountant = data.get('accountant')
    description = data.get('description')

    if expense_type == '其他费用' and not str(description or '').strip():
        return error_response(
            message='费用类型为其他费用时，描述为必填项',
            status_code=status.HTTP_400_BAD_REQUEST,
        )
    if not _project_person_exists(project_id, accountant, '预算员'):
        return error_response(
            message='财务人员必须从当前项目已添加的预算员中选择',
            status_code=status.HTTP_400_BAD_REQUEST,
        )
    serializer = CostInformationSerializer(data=data)
    if serializer.is_valid():
        cost = serializer.save(ID_Project=project)
        create_audit_log(
            request=request,
            action='create',
            module='成本控制',
            target_id=cost.id,
            target_name=project.NAME_Project,
            description=f'创建成本单：{cost.NAME_Cost}，费用类型：{cost.TYPE_Expense}，预算金额：{cost.BUDGET_Amount} {cost.UNIT_Currency}，执行金额：{cost.COST_Amount or "0"}{cost.UNIT_Currency}',
        )
        return success_response(
            data={'cost': CostInformationSerializer(cost).data},
            message='成本单创建成功',
            status_code=status.HTTP_201_CREATED,
        )
    return error_response(
        message='成本单创建失败',
        data=serializer.errors,
        status_code=status.HTTP_400_BAD_REQUEST,
    )

# REST 成本模块：成本单删除接口
@api_view(['PATCH', 'DELETE'])
def rest_cost_detail(request, cost_id):
    try:
        cost = CostInformation.objects.get(id=cost_id)
    except CostInformation.DoesNotExist:
        return error_response(
            message='成本单不存在',
            status_code=status.HTTP_404_NOT_FOUND,
        )
    # 登录校验
    person, auth_error = get_authenticated_person(request)
    if auth_error:
        return auth_error
    # 权限校验
    project = cost.ID_Project
    if not can_manage_project_cost(project, person):
        return permission_denied('只有管理员、当前项目经理或当前项目财务人员可以操作成本单')
    if request.method == 'PATCH':
        data = request.data
        expense_type = data.get('expenseType', cost.TYPE_Expense)
        accountant = data.get('accountant', cost.NAME_Accountant)
        description = data.get('description', cost.DESC_Cost)

        if expense_type == '其他费用' and not str(description or '').strip():
            return error_response(
                message='费用类型为其他费用时，描述为必填项',
                status_code=status.HTTP_400_BAD_REQUEST,
            )
        if accountant and not _project_person_exists(project.id, accountant, '预算员'):
            return error_response(
                message='财务人员必须从当前项目已添加的预算员中选择',
                status_code=status.HTTP_400_BAD_REQUEST,
            )

        serializer = CostInformationSerializer(cost, data=data, partial=True)
        if serializer.is_valid():
            updated_cost = serializer.save()
            create_audit_log(
                request=request,
                action='update',
                module='成本控制',
                target_id=updated_cost.id,
                target_name=project.NAME_Project,
                description=f'修改成本单：{updated_cost.NAME_Cost}，费用类型：{updated_cost.TYPE_Expense}，预算金额：{updated_cost.BUDGET_Amount} {updated_cost.UNIT_Currency}，执行金额：{updated_cost.COST_Amount or "0"}{updated_cost.UNIT_Currency}',
            )
            return success_response(
                data={'cost': CostInformationSerializer(updated_cost).data},
                message='成本单修改成功',
                status_code=status.HTTP_200_OK,
            )
        return error_response(
            message='成本单修改失败',
            data=serializer.errors,
            status_code=status.HTTP_400_BAD_REQUEST,
        )

    try:
        create_audit_log(
            request=request,
            action='delete',
            module='成本控制',
            target_id=cost.pk,
            target_name=project.NAME_Project,
            description=f'删除成本单：{cost.NAME_Cost}，费用类型：{cost.TYPE_Expense}，预算金额：{cost.BUDGET_Amount} {cost.UNIT_Currency}，执行金额：{cost.COST_Amount or "0"}{cost.UNIT_Currency}',
        )
        cost.delete()
        return success_response(
            data=None,
            message='成本单删除成功',
            status_code=status.HTTP_200_OK,
        )
    except Exception as e:
        return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

# ———————————————————— 质量监测 ————————————————————
# REST 质量模板模块：模板列表与新增接口
@api_view(['GET', 'POST'])
def rest_quality_templates(request, project_id):
    if request.method == 'GET':
        try:
            templates = QualityInspectionTemplate.objects.filter(ID_Project=project_id)
            serializer = QualityInspectionTemplateSerializer(templates, many=True)
            return success_response(
                data={'qitTemplates': serializer.data},
                message='获取质检模板列表成功',
                status_code=status.HTTP_200_OK,
            )
        except Exception as e:
            return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

    person, auth_error = get_authenticated_person(request)
    if auth_error:
        return auth_error
    qit_name = request.data.get('qtname')
    items_data = request.data.get('subitems')
    if not project_id or not qit_name or not items_data:
        return error_response(
            message='缺少项目ID、模板名称或检查项',
            status_code=status.HTTP_400_BAD_REQUEST,
        )
    try:
        project = Project.objects.get(id=project_id)
        # 权限校验
        if not can_manage_project_cost(project, person):
            return permission_denied('只有管理员、当前项目经理和当前项目质量员能新建模板')
        qit = QualityInspectionTemplate.objects.create(NAME_Qit=qit_name, ID_Project=project)
        for item_data in items_data:
            InspectionItem.objects.create(
                NAME_Item=item_data.get('name'),
                VALUE_Item=item_data.get('requirement'),
                ID_Qit=qit,
            )
        serializer = QualityInspectionTemplateSerializer(qit)
        create_audit_log(
            request=request,
            action='create',
            module='质量监测',
            target_id=qit.pk,
            target_name=project.NAME_Project,
            description=f'新建模板：{qit.NAME_Qit}',
        )
        return success_response(
            data={'template': serializer.data},
            message='质检模板创建成功',
            status_code=status.HTTP_201_CREATED,
        )
    except Project.DoesNotExist:
        return error_response(message='项目不存在', status_code=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

# REST 质量模板模块：模板删除接口
@api_view(['DELETE'])
def rest_quality_template_detail(request, template_id):
    try:
        person, auth_error = get_authenticated_person(request)
        if auth_error:
            return auth_error
        template = QualityInspectionTemplate.objects.get(id=template_id)
        if not can_manage_project_quality(template.ID_Project, person):
            return permission_denied('只有管理员、当前项目经理和当前项目质量员能删除模板')
        create_audit_log(
            request=request,
            action='delete',
            module='质量监测',
            target_id=template.pk,
            target_name=template.ID_Project.NAME_Project,
            description=f'删除模板：{template.NAME_Qit}',
        )
        template.delete()
        return success_response(
            data=None,
            message='质检模板删除成功',
            status_code=status.HTTP_200_OK,
        )
    except QualityInspectionTemplate.DoesNotExist:
        return error_response(message='质检模板不存在', status_code=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

# REST 质量模板模块：获取模板检查项接口
@api_view(['GET'])
def rest_quality_template_items(request, template_id):
    try:
        items = InspectionItem.objects.filter(ID_Qit=template_id)
        data = [
            {'id': item.pk, 'item_name': item.NAME_Item, 'item_value': item.VALUE_Item}
            for item in items
        ]
        return success_response(
            data={'items': data},
            message='获取质检模板检查项成功',
            status_code=status.HTTP_200_OK,
        )
    except Exception as e:
        return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

# REST 质量报告模块：报告列表与新增接口
@api_view(['GET', 'POST'])
def rest_quality_reports(request, project_id):
    if request.method == 'GET':
        try:
            reports = QualityInspectionReport.objects.filter(ID_Project=project_id)
            serializer = QualityInspectionReportSerializer(reports, many=True)
            return success_response(
                data={'reports': serializer.data},
                message='获取质检报告列表成功',
                status_code=status.HTTP_200_OK,
            )
        except Exception as e:
            return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)
    data = request.data
    try:
        project = Project.objects.get(id=project_id)
        person, auth_error = get_authenticated_person(request)
        if auth_error:
            return auth_error
        if not can_manage_project_cost(project, person):
            return permission_denied('只有管理员、当前项目经理和当前项目质量员能新建报告')
    except Project.DoesNotExist:
        return error_response(message='项目不存在', status_code=status.HTTP_404_NOT_FOUND)

    if not _project_person_exists(project_id, data.get('qrperson'), '质量员'):
        return error_response(
            message='质检员必须从当前项目中的质量员中选择',
            status_code=status.HTTP_400_BAD_REQUEST,
        )
    serializer = QualityInspectionReportSerializer(data=data)
    if not serializer.is_valid():
        return error_response(
            message='质检报告创建失败',
            data=serializer.errors,
            status_code=status.HTTP_400_BAD_REQUEST,
        )
    try:
        report = serializer.save(ID_Project=project)
        for result_data in data.get('qrsubitems', []):
            InspectionResult.objects.create(
                NAME_Item=result_data.get('name'),
                VALUE_Standard=result_data.get('requirement'),
                RESULT_Inspect=result_data.get('result'),
                ID_Report=report,
            )
        create_audit_log(
            request=request,
            action='create',
            module='质量监测',
            target_id=report.id,
            target_name=project.NAME_Project,
            description=f'新增质量报告：工程名称：{report.NAME_Project}，检验部位及编号：{report.PART_Num}，总体情况：{report.STATUS_Inspect}',
        )
        return success_response(
            data={'report': QualityInspectionReportSerializer(report).data},
            message='质检报告创建成功',
            status_code=status.HTTP_201_CREATED,
        )
    except Exception as e:
        return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

# REST 质量报告模块：报告详情、修改与删除接口
@api_view(['GET', 'PATCH', 'DELETE'])
def rest_quality_report_detail(request, report_id):
    try:
        report = QualityInspectionReport.objects.get(id=report_id)
    except QualityInspectionReport.DoesNotExist:
        return error_response(message='质量报告不存在', status_code=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        report_data = QualityInspectionReportSerializer(report).data
        report_data['qrsubitems'] = [
            {
                'name': result.NAME_Item,
                'requirement': result.VALUE_Standard,
                'result': result.RESULT_Inspect,
            }
            for result in InspectionResult.objects.filter(ID_Report=report)
        ]
        return success_response(
            data={'report': report_data},
            message='获取质量报告详情成功',
            status_code=status.HTTP_200_OK,
        )
    person, auth_error = get_authenticated_person(request)
    if auth_error:
        return auth_error

    project = report.ID_Project
    if project is None:
        return error_response(
            message='质量报告所属项目不存在',
            status_code=status.HTTP_400_BAD_REQUEST,
        )
    if not can_manage_project_quality(project, person):
        action_text = '修改' if request.method == 'PATCH' else '删除'
        return permission_denied(f'只有管理员、当前项目经理或当前项目质量员可以{action_text}质量报告')

    if request.method == 'PATCH':
        data = request.data
        if not _project_person_exists(project.id, data.get('qrperson'), '质量员'):
            return error_response(
                message='质检员必须从当前项目中的质量员中选择',
                status_code=status.HTTP_400_BAD_REQUEST,
            )
        serializer = QualityInspectionReportSerializer(report, data=data, partial=True)
        if not serializer.is_valid():
            return error_response(
                message='质量报告修改失败',
                data=serializer.errors,
                status_code=status.HTTP_400_BAD_REQUEST,
            )
        report = serializer.save()
        if 'qrsubitems' in data:
            InspectionResult.objects.filter(ID_Report=report).delete()
            for result_data in data.get('qrsubitems', []):
                InspectionResult.objects.create(
                    NAME_Item=result_data.get('name'),
                    VALUE_Standard=result_data.get('requirement'),
                    RESULT_Inspect=result_data.get('result'),
                    ID_Report=report,
                )
        create_audit_log(
            request=request,
            action='update',
            module='质量监测',
            target_id=report.id,
            target_name=project.NAME_Project,
            description=f'修改质量报告：工程名称：{report.NAME_Project}，检验部位及编号：{report.PART_Num}，总体情况：{report.STATUS_Inspect}',
        )
        return success_response(
            data={'report': QualityInspectionReportSerializer(report).data},
            message='质量报告修改成功',
            status_code=status.HTTP_200_OK,
        )
    create_audit_log(
        request=request,
        action='delete',
        module='质量监测',
        target_id=report.pk,
        target_name=project.NAME_Project,
        description=f'删除质量报告：工程名称：{report.NAME_Project}，检验部位及编号：{report.PART_Num}，总体情况：{report.STATUS_Inspect}',
    )
    report.delete()
    return success_response(
        data=None,
        message='质量报告删除成功',
        status_code=status.HTTP_200_OK,
    )

# 质量监测统计图
@api_view(['GET'])
def rest_quality_stats(request, project_id):
    try:
        stats = {
            '合格': [0, 0, 0, 0],
            '一般质量问题': [0, 0, 0, 0],
            '重大质量问题': [0, 0, 0, 0],
        }
        report_stats = (
            QualityInspectionReport.objects.filter(ID_Project=project_id)
            .annotate(quarter=ExtractQuarter('TIME_Inspect'))
            .values('STATUS_Inspect', 'quarter')
            .annotate(count=Count('STATUS_Inspect'))
        )
        for stat in report_stats:
            status_value = stat['STATUS_Inspect']
            quarter = stat['quarter']
            count = stat['count']
            if status_value in stats and quarter:
                stats[status_value][quarter - 1] = count
        return success_response(
            data={'stats': stats},
            message='获取质量问题统计成功',
            status_code=status.HTTP_200_OK,
        )
    except Exception as e:
        return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

# ———————————————————— 安全监测 ————————————————————
# 获取安全检查项
@api_view(['GET'])
def sitem_list(request, scr_id):
    try:
        items = SecurityCheckItem.objects.filter(ID_Sct=scr_id)
        items_data = [
            {
                'id': item.pk,
                'item_name': item.NAME_Item,
                'item_value': item.VALUE_Item,
            }
            for item in items
        ]
        return success_response(
            data={'items': items_data},
            message='获取安全检查项成功',
            status_code=status.HTTP_200_OK,
        )
    except Exception as e:
        return error_response(
            message=str(e),
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )

# REST 安全模板模块：模板列表与新增接口
@api_view(['GET', 'POST'])
def rest_safety_templates(request, project_id):
    if request.method == 'GET':
        try:
            templates = SecurityCheckTemplate.objects.filter(ID_Project=project_id)
            serializer = SecurityCheckTemplateSerializer(templates, many=True)
            return success_response(
                data={'sctTemplates': serializer.data},
                message='获取安全模板列表成功',
                status_code=status.HTTP_200_OK,
            )
        except Exception as e:
            return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

    template_name = request.data.get('stname')
    items_data = request.data.get('subitems')
    # 登录校验
    person, auth_error = get_authenticated_person(request)
    if auth_error:
        return auth_error
    # 权限校验
    project = Project.objects.get(id=project_id)
    if not can_manage_project_safety(project, person):
        return permission_denied('只有管理员、当前项目经理或当前项目安全人员可以操作')

    if not project_id or not template_name or not items_data:
        return error_response(
            message='缺少项目ID、模板名称或检查项',
            status_code=status.HTTP_400_BAD_REQUEST,
        )
    try:
        project = Project.objects.get(id=project_id)
        template = SecurityCheckTemplate.objects.create(NAME_Sct=template_name, ID_Project=project)
        for item_data in items_data:
            SecurityCheckItem.objects.create(
                NAME_Item=item_data.get('name'),
                VALUE_Item=item_data.get('requirement'),
                ID_Sct=template,
            )
        serializer = SecurityCheckTemplateSerializer(template)
        create_audit_log(
            request=request,
            action='create',
            module='安全监测',
            target_id=template.pk,
            target_name=project.NAME_Project,
            description=f'新增安全模板：{template.NAME_Sct}',
        )
        return success_response(
            data={'template': serializer.data},
            message='安全模板创建成功',
            status_code=status.HTTP_201_CREATED,
        )
    except Project.DoesNotExist:
        return error_response(message='项目不存在', status_code=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

# REST 安全模板模块：模板删除接口
@api_view(['DELETE'])
def rest_safety_template_detail(request, template_id):
    try:
        template = SecurityCheckTemplate.objects.get(id=template_id)
        # 登录校验
        person, auth_error = get_authenticated_person(request)
        if auth_error:
            return auth_error
        # 权限校验
        project = template.ID_Project
        if not can_manage_project_safety(project, person):
            return permission_denied('只有管理员、当前项目经理或当前项目安全人员可以操作')
        create_audit_log(
            request=request,
            action='delete',
            module='安全监测',
            target_id=template.pk,
            target_name=project.NAME_Project,
            description=f'删除安全模板：{template.NAME_Sct}',
        )
        template.delete()
        return success_response(
            data=None,
            message='安全模板删除成功',
            status_code=status.HTTP_200_OK,
        )
    except SecurityCheckTemplate.DoesNotExist:
        return error_response(message='安全模板不存在', status_code=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

# REST 安全模板模块：获取模板检查项接口
@api_view(['GET'])
def rest_safety_template_items(request, template_id):
    try:
        items = SecurityCheckItem.objects.filter(ID_Sct=template_id)
        data = [
            {'id': item.pk, 'item_name': item.NAME_Item, 'item_value': item.VALUE_Item}
            for item in items
        ]
        return success_response(
            data={'items': data},
            message='获取安全模板检查项成功',
            status_code=status.HTTP_200_OK,
        )
    except Exception as e:
        return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

# 新建安全检验报告
@api_view(['POST'])
def rest_safety_reports(request, project_id):
    report_data = request.POST.get('report')
    if not report_data:
        return error_response(
            message='缺少报告数据',
            status_code=status.HTTP_400_BAD_REQUEST,
        )
    try:
        data = json.loads(report_data)
        images = request.FILES.getlist('images')
        project = Project.objects.get(id=project_id)
    except Project.DoesNotExist:
        return error_response(message='项目不存在', status_code=status.HTTP_404_NOT_FOUND)
    except json.JSONDecodeError:
        return error_response(message='报告数据格式错误', status_code=status.HTTP_400_BAD_REQUEST)
    except Exception as e:
        return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)
    # 登录校验
    person, auth_error = get_authenticated_person(request)
    if auth_error:
        return auth_error
    # 权限校验
    if not can_manage_project_safety(project, person):
        return permission_denied('只有管理员、当前项目经理或当前项目安全人员可以新建安全报告')

    if not _project_person_exists(project_id, data.get('srperson'), '安全员'):
        return error_response(
            message='安全员必须从当前项目中的安全员中选择',
            status_code=status.HTTP_400_BAD_REQUEST,
        )
    serializer = SecurityInspectionReportSerializer(data=data)
    if not serializer.is_valid():
        return error_response(
            message='安全报告创建失败',
            data=serializer.errors,
            status_code=status.HTTP_400_BAD_REQUEST,
        )
    try:
        report = serializer.save(ID_Project=project)
        for item in data.get('srsubitems', []):
            SecurityCheckResult.objects.create(
                NAME_Item=item.get('name'),
                STANDARD_Item=item.get('requirement'),
                RESULT_Item=item.get('result'),
                ID_Report=report,
            )
        for image in images:
            SecurityInspectionImage.objects.create(image=image, ID_Report=report)
        create_audit_log(
            request=request,
            action='create',
            module='安全监测',
            target_id=report.pk,
            target_name=project.NAME_Project,
            description=f'新增安全报告：工程名称：{report.NAME_Project}，检查部位及编号：{report.PART_Check}，报告编号：{report.NUM_Report}，总体情况：{report.STATUS_Overall}',
        )
        return success_response(
            data={'report': SecurityInspectionReportSerializer(report).data},
            message='安全报告创建成功',
            status_code=status.HTTP_201_CREATED,
        )
    except Exception as e:
        return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

# 现存安全问题列表
@api_view(['GET'])
def rest_safety_issues(request, project_id):
    try:
        reports = SecurityInspectionReport.objects.filter(
            ID_Project=project_id,
            STATUS_Overall__in=['一般安全问题', '重大安全问题'],
            solutions__isnull=True,
        ).distinct().order_by('-id')
        serializer = SecurityInspectionReportSerializer(reports, many=True)
        return success_response(
            data={'filteredReports': serializer.data},
            message='获取现存安全问题成功',
            status_code=status.HTTP_200_OK,
        )
    except Exception as e:
        return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

# REST 安全问题模块：删除现存安全问题接口
@api_view(['DELETE'])
def rest_safety_issue_detail(request, report_id):
    try:
        report = SecurityInspectionReport.objects.get(
            id=report_id,
            STATUS_Overall__in=['一般安全问题', '重大安全问题'],
            solutions__isnull=True,
        )
    except SecurityInspectionReport.DoesNotExist:
        return error_response(message='现存安全问题不存在', status_code=status.HTTP_404_NOT_FOUND)
    # 登录校验
    person, auth_error = get_authenticated_person(request)
    if auth_error:
        return auth_error
    # 权限校验
    if not can_manage_project_safety(report.ID_Project, person):
        return permission_denied('只有管理员、当前项目经理或当前项目安全人员可以操作')
    try:
        create_audit_log(
            request=request,
            action='delete',
            module='安全监测',
            target_id=report.pk,
            target_name=report.ID_Project.NAME_Project,
            description=f'删除安全问题：工程名称：{report.NAME_Project}，检查部位及编号：{report.PART_Check}，报告编号：{report.NUM_Report}，总体情况：{report.STATUS_Overall}',
        )
        report.delete()
        return success_response(
            data=None,
            message='现存安全问题删除成功',
            status_code=status.HTTP_200_OK,
        )
    except Exception as e:
        return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

# REST 已处理安全问题列表接口
@api_view(['GET'])
def rest_safety_resolved_issues(request, project_id):
    try:
        reports = SecurityInspectionReport.objects.filter(
            ID_Project=project_id,
            solutions__isnull=False,
        ).distinct().order_by('-id')
        report_number = request.query_params.get('reportNumber') or request.query_params.get('srnumber')
        if report_number:
            reports = reports.filter(NUM_Report__icontains=report_number)
        serializer = SecurityInspectionReportSerializer(reports, many=True)
        return success_response(
            data={'resolvedReports': serializer.data},
            message='获取已处理安全问题成功',
            status_code=status.HTTP_200_OK,
        )
    except Exception as e:
        return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

# 处理安全问题
@api_view(['GET', 'POST'])
def rest_safety_report_solutions(request, report_id):
    if request.method == 'GET':
        try:
            solutions = SafetyIssueSolution.objects.filter(ID_Report=report_id)
            if not solutions.exists():
                return error_response(
                    message='未找到安全问题处理记录',
                    status_code=status.HTTP_404_NOT_FOUND,
                )
            serializer = SafetyIssueSolutionSerializer(solutions, many=True)
            return success_response(
                data={'solutions': serializer.data},
                message='获取安全问题处理记录成功',
                status_code=status.HTTP_200_OK,
            )
        except Exception as e:
            return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

    solution_date = request.data.get('res_Date') or request.data.get('date')
    solution_description = request.data.get('resolution') or request.data.get('description')
    if not report_id or not solution_date or not solution_description:
        return error_response(
            message='请填写完整的处理信息',
            status_code=status.HTTP_400_BAD_REQUEST,
        )
    try:
        report = SecurityInspectionReport.objects.get(id=report_id)
        # 登录校验
        person, auth_error = get_authenticated_person(request)
        if auth_error:
            return auth_error
        # 权限校验
        if not can_manage_project_safety(report.ID_Project, person):
            return permission_denied('只有管理员、当前项目经理或当前项目安全人员可以操作')
        SafetyIssueSolution.objects.create(
            ID_Report=report,
            DATE_Solution=solution_date,
            DESCRIPTION_Solution=solution_description,
        )
        report.STATUS_Overall = '已处理'
        report.save()
        create_audit_log(
            request=request,
            action='update',
            module='安全监测',
            target_id=report.pk,
            target_name=report.ID_Project.NAME_Project,
            description=f'处理安全问题：工程名称：{report.NAME_Project}，检查部位及编号：{report.PART_Check}，报告编号：{report.NUM_Report}，总体情况：{report.STATUS_Overall}',
        )
        return success_response(
            data=None,
            message='安全问题处理成功',
            status_code=status.HTTP_201_CREATED,
        )
    except SecurityInspectionReport.DoesNotExist:
        return error_response(message='安全报告不存在', status_code=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

# ———————————————————— 文件 ————————————————————
# REST 文件模块：项目文件列表与上传接口
@api_view(['GET', 'POST'])
def rest_project_files(request, project_id):
    if request.method == 'GET':
        try:
            files = File.objects.filter(ID_Project=project_id)
            serializer = FileSerializer(files, many=True)
            files_data = list(serializer.data)
            try:
                indexed_file_ids = get_indexed_file_ids(project_id)
            except Exception:
                indexed_file_ids = set()
            for file_data in files_data:
                file_id = file_data.get('file_id')
                file_data['is_indexed'] = file_id in indexed_file_ids
            return success_response(
                data={'files': files_data},
                message='获取文件列表成功',
                status_code=status.HTTP_200_OK,
            )
        except Exception as e:
            return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)
    # 登录校验
    person, auth_error = get_authenticated_person(request)
    if auth_error:
        return auth_error

    if not request.FILES.get('file'):
        return error_response(
            message='请选择要上传的文件',
            status_code=status.HTTP_400_BAD_REQUEST,
        )
    try:
        uploaded_file = request.FILES['file']
        file_name, file_extension = os.path.splitext(uploaded_file.name)
        project_instance = Project.objects.get(id=project_id)
        # 权限校验
        if not can_manage_project_file(project_instance, person):
            return permission_denied('只有管理员、当前项目经理或当前项目资料人员可以上传文件')
        file_instance = File.objects.create(
            FILE=uploaded_file,
            NAME_File=file_name,
            SIZE_File=uploaded_file.size,
            FORM_File=file_extension,
            ID_Project=project_instance,
            UPLOADER_Person=person,
        )
        create_audit_log(
            request=request,
            action='upload',
            module='文件',
            target_id=file_instance.pk,
            target_name=project_instance.NAME_Project,
            description=f'上传文件：{file_instance.NAME_File}{file_instance.FORM_File or ""}',
        )
        return success_response(
            data={'file': FileSerializer(file_instance).data},
            message='文件上传成功',
            status_code=status.HTTP_201_CREATED,
        )
    except Project.DoesNotExist:
        return error_response(
            message='项目不存在',
            status_code=status.HTTP_404_NOT_FOUND,
        )
    except Exception as e:
        return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

# 文件预览
@api_view(['GET'])
def file_preview(request, file_id=None):
    try:
        file_obj = File.objects.get(id=file_id)
        person, auth_error = get_authenticated_person(request)
        if auth_error:
            html_content = (
                '<html><head><meta charset="utf-8"><style>'
                'body{margin:0;height:100vh;display:flex;align-items:center;justify-content:center;font-family:Segoe UI,Arial,sans-serif;}'
                '.error{color:#d32f2f;font-size:20px;}'
                '</style></head><body>'
                '<div class="error">请先登录</div>'
                '</body></html>'
            )
            return HttpResponse(html_content, status=401, content_type='text/html; charset=utf-8')
        # 权限校验
        if not is_admin(person) and not is_project_member(file_obj.ID_Project, person):
            html_content = (
                '<html><head><meta charset="utf-8"><style>'
                'body{margin:0;height:100vh;display:flex;align-items:center;justify-content:center;font-family:Segoe UI,Arial,sans-serif;}'
                '.error{color:#d32f2f;font-size:20px;}'
                '</style></head><body>'
                '<div class="error">只有当前项目成员可以预览文件</div>'
                '</body></html>'
            )
            return HttpResponse(html_content, status=403, content_type='text/html; charset=utf-8')
        file_path = file_obj.FILE.path
        file_extension = (file_obj.FORM_File or '').lower()
        file_name = f"{file_obj.NAME_File}{file_obj.FORM_File or ''}"

        if file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']:
            with open(file_path, 'rb') as f:
                file_content = f.read()
            mime_type = mimetypes.guess_type(file_name)[0] or 'application/octet-stream'
            response = HttpResponse(file_content, content_type=mime_type)
            response['Content-Disposition'] = f'inline; filename="{smart_str(file_name)}"'
            create_audit_log(
                request=request,
                action='preview',
                module='文件',
                target_id=file_obj.pk,
                target_name=file_obj.ID_Project.NAME_Project,
                description=f'预览文件：文件名：{file_name}',
            )
            return response

        if file_extension == '.pdf':
            with open(file_path, 'rb') as f:
                file_content = f.read()
            response = HttpResponse(file_content, content_type='application/pdf')
            response['Content-Disposition'] = f'inline; filename="{smart_str(file_name)}"'
            create_audit_log(
                request=request,
                action='preview',
                module='文件',
                target_id=file_obj.pk,
                target_name=file_obj.ID_Project.NAME_Project,
                description=f'预览文件：文件名：{file_name}',
            )
            return response

        if file_extension in ['.txt', '.md', '.py', '.json', '.js', '.css', '.html']:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                text_content = f.read()
            html_content = (
                '<html><head><meta charset="utf-8"><style>'
                'body{font-family:Segoe UI,Arial,sans-serif;padding:24px;line-height:1.7;white-space:pre-wrap;word-break:break-word;}'
                '</style></head><body>'
                f'{escape(text_content)}'
                '</body></html>'
            )
            create_audit_log(
                request=request,
                action='preview',
                module='文件',
                target_id=file_obj.pk,
                target_name=file_obj.ID_Project.NAME_Project,
                description=f'预览文件：文件名：{file_name}',
            )
            return HttpResponse(html_content, content_type='text/html; charset=utf-8')

        if file_extension == '.docx':
            document = Document(file_path)
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            table_rows = []
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        table_rows.append(' | '.join(cells))
            combined_lines = paragraphs + table_rows
            preview_text = '\n\n'.join(combined_lines) if combined_lines else '文档中暂无可预览的文字内容'
            html_content = (
                '<html><head><meta charset="utf-8"><style>'
                'body{font-family:Segoe UI,Arial,sans-serif;padding:24px;line-height:1.8;white-space:pre-wrap;word-break:break-word;}'
                'h2{margin-top:0;} .meta{color:#666;margin-bottom:16px;font-size:14px;}'
                '</style></head><body>'
                f'<h2>{escape(file_name)}</h2>'
                f'<div class="meta">{escape("Word 预览")}</div>'
                f'{escape(preview_text)}'
                '</body></html>'
            )
            create_audit_log(
                request=request,
                action='preview',
                module='文件',
                target_id=file_obj.pk,
                target_name=file_obj.ID_Project.NAME_Project,
                description=f'预览文件：文件名：{file_name}',
            )
            return HttpResponse(html_content, content_type='text/html; charset=utf-8')

        if file_extension in ['.xlsx', '.xlsm']:
            workbook = load_workbook(file_path, data_only=True)
            sheets_html = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                rows_html = []

                for row in sheet.iter_rows(values_only=True):
                    cells_html = ''.join(
                        f'<td>{escape(str(cell)) if cell is not None else ""}</td>'
                        for cell in row
                    )
                    rows_html.append(f'<tr>{cells_html}</tr>')

                sheets_html.append(
                    f'<h3>{escape(sheet_name)}</h3>'
                    f'<table>{"".join(rows_html)}</table>'
                )
            html_content = (
                '<html><head><meta charset="utf-8"><style>'
                'body{font-family:Segoe UI,Arial,sans-serif;padding:24px;line-height:1.7;}'
                'h2{margin-top:0;} h3{margin-top:24px;}'
                'table{border-collapse:collapse;margin-bottom:24px;max-width:100%;}'
                'td{border:1px solid #ddd;padding:6px 10px;min-width:80px;}'
                '</style></head><body>'
                f'<h2>{escape(file_name)}</h2>'
                f'{"".join(sheets_html)}'
                '</body></html>'
            )
            create_audit_log(
                request=request,
                action='preview',
                module='文件',
                target_id=file_obj.pk,
                target_name=file_obj.ID_Project.NAME_Project,
                description=f'预览文件：文件名：{file_name}',
            )
            return HttpResponse(html_content, content_type='text/html; charset=utf-8')

        html_content = (
            '<html><head><meta charset="utf-8"><style>'
            'body{font-family:Segoe UI,Arial,sans-serif;padding:24px;line-height:1.7;}'
            '.note{padding:16px;border-radius:8px;background:#f5f7fb;color:#333;}'
            '</style></head><body>'
            f'<h2>{escape(file_name)}</h2>'
            '<div class="note">暂不支持当前文件格式的在线预览，请下载后在本地打开查看。</div>'
            '</body></html>'
        )
        create_audit_log(
            request=request,
            action='preview',
            module='文件',
            target_id=file_obj.pk,
            target_name=file_obj.ID_Project.NAME_Project,
            description=f'预览文件：文件名：{file_name}',
        )
        return HttpResponse(html_content, content_type='text/html; charset=utf-8')
    except File.DoesNotExist:
        return JsonResponse({'error': 'File not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

# 文件下载
@api_view(['GET'])
def file_download(request, file_id=None):
    # 登录校验
    person, auth_error = get_authenticated_person(request)
    if auth_error:
        return auth_error
    try:
        file_obj = File.objects.get(id=file_id)
        file_path = file_obj.FILE.path
        # 权限校验
        if not is_admin(person) and not is_project_member(file_obj.ID_Project, person):
            return permission_denied('只有当前项目成员可以下载文件')

        with open(file_path, 'rb') as f:
            file_content = f.read()

        response = HttpResponse(file_content, content_type='application/octet-stream; charset=utf-8')
        response['Content-Disposition'] = f'attachment; filename="{smart_str(file_obj.NAME_File)}"'
        response['Content-Length'] = os.path.getsize(file_path)
        create_audit_log(
            request=request,
            action='download',
            module='文件',
            target_id=file_obj.pk,
            target_name=file_obj.ID_Project.NAME_Project,
            description=f'下载文件：文件名：{file_obj.NAME_File}{file_obj.FORM_File or ""}',
        )
        return response
    except File.DoesNotExist:
        return JsonResponse({'error': 'File not found'}, status=404)
    except Exception as e:
        print(f'Exception: {e}')
        return JsonResponse({'error': str(e)}, status=500)

# REST 文件模块：文件删除接口
@api_view(['DELETE'])
def rest_files(request, file_id):
    # 登录校验
    person, auth_error = get_authenticated_person(request)
    if auth_error:
        return auth_error
    try:
        file_obj = File.objects.get(id=file_id)
        # 权限校验
        if not can_manage_project_file(file_obj.ID_Project, person):
            return permission_denied('只有管理员、当前项目经理或当前项目资料人员可以删除文件')
        create_audit_log(
            request=request,
            action='delete',
            module='文件',
            target_id=file_obj.pk,
            target_name=file_obj.ID_Project.NAME_Project,
            description=f'删除文件：文件名：{file_obj.NAME_File}{file_obj.FORM_File or ""}',
        )
        # 从 ES 和 Qdrant 中同步删除
        delete_langchain_file_vectors(file_obj.pk)
        delete_file_chunks_from_elasticsearch(file_obj.pk)
        file_obj.delete()
        return success_response(
            data=None,
            message='文件删除成功',
            status_code=status.HTTP_200_OK,
        )
    except File.DoesNotExist:
        return error_response(
            message='文件不存在',
            status_code=status.HTTP_404_NOT_FOUND,
        )
    except Exception as e:
        return error_response(message=str(e), status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

# 文件入 Qdrant 库
@api_view(['POST'])
def rag_index_file(request, file_id):
    person, auth_error = get_authenticated_person(request)
    if auth_error:
        return auth_error

    try:
        file_obj = File.objects.get(id=file_id)
    except File.DoesNotExist:
        return error_response(
            message='文件不存在',
            status_code=status.HTTP_404_NOT_FOUND,
        )

    if not can_manage_project_file(file_obj.ID_Project, person):
        return permission_denied('只有管理员、当前项目经理或当前项目资料人员可以进行文件入库')

    try:
        result = index_file_to_qdrant_langchain(file_obj)
        return success_response(
            data=result,
            message='文件向量入库成功',
            status_code=status.HTTP_200_OK,
        )
    except Exception as e:
        return error_response(
            message=str(e),
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )


# 文件重新入库
@api_view(['POST'])
def rag_reindex_file(request, file_id):
    person, auth_error = get_authenticated_person(request)
    if auth_error:
        return auth_error
    try:
        file_obj = File.objects.get(id=file_id)
    except File.DoesNotExist:
        return error_response(
            message='文件不存在',
            status_code=status.HTTP_404_NOT_FOUND,
        )
    project = file_obj.ID_Project
    if not can_manage_project_file(project, person):
        return permission_denied('只有管理员、当前项目经理或当前项目资料员可以重新入库文件')
    try:
        result = index_file_to_qdrant_langchain(file_obj)
        return success_response(
            data=result,
            message='文件重新入库成功',
            status_code=status.HTTP_200_OK,
        )
    except Exception as e:
        return error_response(
            message=str(e),
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )

# 删除文件向量
@api_view(['DELETE'])
def rag_delete_file_vectors(request, file_id):
    person, auth_error = get_authenticated_person(request)
    if auth_error:
        return auth_error

    try:
        file_obj = File.objects.get(id=file_id)
    except File.DoesNotExist:
        return error_response(
            message='文件不存在',
            status_code=status.HTTP_404_NOT_FOUND,
        )

    project = file_obj.ID_Project

    if not can_manage_project_file(project, person):
        return permission_denied('只有管理员、当前项目经理或当前项目资料员可以删除文件向量')

    try:
        delete_langchain_file_vectors(file_obj.pk)
        return success_response(
            data=None,
            message='文件向量删除成功',
            status_code=status.HTTP_200_OK,
        )
    except Exception as e:
        return error_response(
            message=str(e),
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )


# 项目问答接口
@api_view(['POST'])
def rag_chat(request, project_id):
    person, auth_error = get_authenticated_person(request)
    if auth_error:
        return auth_error

    try:
        project = Project.objects.get(id=project_id)
    except Project.DoesNotExist:
        return error_response(
            message='项目不存在',
            status_code=status.HTTP_404_NOT_FOUND,
        )

    if not is_admin(person) and not is_project_member(project, person):
        return permission_denied('只有当前项目成员可以使用项目文档问答')

    question = request.data.get('question')
    history = request.data.get('history') or []
    if not question:
        return error_response(
            message='请输入问题',
            status_code=status.HTTP_400_BAD_REQUEST,
        )


    def stream_events():
        try:
            for event in answer_question_with_rag(
                question=question,
                project_id=project_id,
                limit=8,
                history=history,
            ):
                yield json.dumps(event, ensure_ascii=False) + '\n'
        except Exception as e:
            yield json.dumps(
                {
                    'type': 'error',
                    'message': str(e),
                },
                ensure_ascii=False,
            ) + '\n'

    response = StreamingHttpResponse(
        stream_events(),
        content_type='application/x-ndjson; charset=utf-8',
    )
    response['Cache-Control'] = 'no-cache'
    return response
